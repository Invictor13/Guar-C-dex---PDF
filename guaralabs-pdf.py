import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from pypdf import PdfReader, PdfWriter
import fitz # PyMuPDF
from docx import Document
import pandas as pd
import os
import tempfile
from io import BytesIO
from PIL import Image, ImageTk
import math
import random
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import logging

# Configurar logging
logging.basicConfig(filename="guara_codex.log", level=logging.INFO,
                    format="%(asctime)s - %(levelname)s - %(module)s - %(funcName)s - %(message)s")

# --- Textos da Interface (Diretamente em Português) ---
texts = {
    "title": "Guará Codex - Editor de PDF",
    "manipulation_frame": "Manipulação de Páginas 🐺",
    "conversion_frame": "Conversão e Otimização ⚙️",
    "split_pdf": "Separar PDF 🪚",
    "merge_pdf": "Mesclar PDF 📚",
    "exclude_pages": "Excluir Página(s) 🗑️",
    "select_pages": "Selecionar Página(s) 📜",
    "add_pages": "Acrescentar Página(s) 📋",
    "replace_pages": "Substituir Página(s) 🔄",
    "compress_pdf": "Compactar PDF 🗃️",
    "anonymize_pdf": "Anonimizar PDF (LGPD) 🔒",
    "convert_pdf": "Converter PDF 📄",
    "pdf_jpg": "PDF ↔️ JPG 🖼️",
    "exit": "Sair 🚪",
    "help": "Ajuda ❓",
    "choose_conversion": "Escolha o tipo de conversão:",
    "pdf_to_jpg": "PDF para JPG",
    "jpg_to_pdf": "JPG para PDF",
    "choose_output": "Selecione o formato de saída:",
    "word": "Word (.docx)",
    "excel": "Excel (.xlsx)",
    "split_option": "Selecione a opção de divisão:",
    "split_half": "Dividir pela metade",
    "split_third": "Dividir em 1/3",
    "split_quarter": "Dividir em 1/4",
    "split_custom": "Customizar",
    "custom_parts_prompt": "Digite o número de partes (ex.: 5):",
    "interactive_exclude_prompt": "Páginas a Excluir (ex: 1-3 ou 1,3,5):",
    "interactive_select_prompt": "Páginas a Selecionar (ex: 1-3 ou 1,3,5):",
    "interactive_add_prompt": "Acrescentar APÓS a página (0 para início, {total_original} para fim):",
    "interactive_replace_prompt": "Páginas a Substituir no Original (ex: 1-3 ou 5):",
    "interactive_file_original_label": "Arquivo PDF Original:",
    "interactive_file_second_label": "Arquivo PDF Adicional/Substituição:",
    "interactive_page_range_description_select": "Digite o intervalo ou páginas específicas a serem extraídas. A prévia mostrará as páginas selecionadas.",
    "interactive_page_range_description_exclude": "Digite o intervalo ou páginas específicas a serem removidas. A prévia mostrará as páginas a serem excluídas.",
    "interactive_add_description": "O PDF Adicional será inserido após a página indicada do PDF Original.",
    "interactive_replace_description": "As páginas do PDF Adicional substituirão sequencialmente o intervalo indicado no PDF Original.",
    "compress_level": "Selecione o nível de compactação:",
    "light": "Leve (Alta Qualidade)",
    "moderate": "Moderada (Média Qualidade)",
    "aggressive": "Agressiva (Baixa Qualidade)",
    "success": "Sucesso",
    "error": "Erro",
    "footer": "Desenvolvido por GuaráLabs",
    "tooltip_split": "Divida PDFs em partes iguais com precisão!",
    "tooltip_merge": "Combine vários PDFs em um único arquivo!",
    "tooltip_exclude": "Remova páginas específicas do PDF rapidamente com prévia interativa!",
    "tooltip_select": "Extraia páginas específicas do PDF com facilidade com prévia interativa!",
    "tooltip_add": "Insira páginas de outro PDF na posição desejada com prévia interativa!",
    "tooltip_replace": "Substitua páginas do PDF de forma eficiente com prévia interativa!",
    "tooltip_compress": "Reduza o tamanho do PDF sem perder qualidade (sem OCR)!",
    "tooltip_anonymize": "Anonimize PDFs e garanta conformidade LGPD em segundos!",
    "tooltip_convert": "Converta PDFs para Word ou Excel em um clique!",
    "tooltip_pdf_jpg": "Transforme PDFs em JPG ou vice-versa rapidamente!",
    "tooltip_exit": "Encerre sua caçada com segurança!",
    "tooltip_help": "Descubra todas as funções do Guará Codex!",
    "help_text": (
        "Bem-vindo ao Guará Codex - Editor de PDF!\n\n"
        "Funções:\n"
        "- Separar PDF: Divide em partes iguais.\n"
        "- Mesclar PDF: Combina PDFs.\n"
        "- Excluir Página(s): Abre janela interativa para definir e pré-visualizar páginas a remover.\n"
        "- Selecionar Página(s): Abre janela interativa para definir e pré-visualizar páginas a extrair.\n"
        "- Acrescentar Página(s): Abre janela interativa para selecionar PDFs, definir ponto de inserção e pré-visualizar.\n"
        "- Substituir Página(s): Abre janela interativa para selecionar PDFs, definir intervalo e pré-visualizar substituição.\n"
        "- Compactar PDF: Reduz o tamanho do PDF (renderizando como imagem).\n"
        "- Anonimizar PDF: Remove metadados (LGPD).\n"
        "- Converter PDF: Converte PDF para Word ou Excel.\n"
        "- PDF ↔️ JPG: Converte entre formatos.\n\n"
        "Atalhos: Ctrl+C (compactar), Esc (sair).\n"
        "Desenvolvido por GuaráLabs."
    ),
    "invalid_range": "Ops, intervalo inválido! 🐺 O PDF tem {total} páginas. Tente algo como 1-3 ou 1,3,5.",
    "invalid_parts": "Ops, número de partes inválido! 🐺 Deve ser um número inteiro maior que 1.",
    "invalid_insert_point": "Ponto de inserção inválido. Deve ser entre 0 e {total}.",
    "processing_page": "Processando página {current} de {total}... (ETA: {eta:.1f}s)",
    "compress_success": "PDF compactado com sucesso! 🐺 Tamanho reduzido de {old_size:.2f}MB para {new_size:.2f}MB.",
    "preview_title_select": "Selecionar Páginas",
    "preview_title_exclude": "Excluir Páginas",
    "preview_title_add": "Acrescentar Páginas",
    "preview_title_replace": "Substituir Páginas",
    "preview_confirm": "Confirmar Operação",
    "preview_cancel": "Cancelar",
    "preview_update_button": "Atualizar Prévia",
    "preview_page_label": "Página {num}",
    "preview_page_of_label": "Página {num} de {pdf_name}",
    "preview_add_insert_label": "-> Págs. do 2º PDF aqui <-",
    "preview_replace_original_label": "Original Pág. {num}",
    "preview_replace_new_label": "Nova Pág. {num}",
    "preview_info_original_pdf": "PDF Original: {filename}",
    "preview_info_second_pdf": "PDF Adicional/Substituição: {filename}",
}

# --- Constantes Globais e Variáveis (para a UI) ---
preview_images = [] # Lista global para manter referências de imagens Tkinter
MAX_PREVIEW_IMAGES_SHOWN = 10 # Limite de imagens mostradas na prévia para performance
PREVIEW_RESOLUTION_FACTOR = 1.5 # Fator de resolução para renderização de prévias
MAX_PREVIEW_IMG_WIDTH_SINGLE = 450 # Largura máxima para imagens de prévia em painel único
MAX_PREVIEW_IMG_WIDTH_DUAL = 380 # Largura máxima para imagens de prévia em painéis duplos

# Definindo as cores da Toca para fácil referência
COLOR_TERRACOTTA = '#e67e22' # Terracota do Cerrado (Primária da Alcateia)
COLOR_NIGHT_SKY = '#2c2c2c' # Noite Estrelada (Fundo Escuro Principal para contraste)
COLOR_MOON_LIGHT = '#ffffff' # Luz Pura da Lua (Fundo Claro Principal / Textos CLAROS)
COLOR_DARK_EARTH = '#4a3726' # Um marrom escuro para textos e detalhes, baseado no contraste.
COLOR_ACCENT_GREEN = '#6B8E23' # Um verde musgo ou cerrado para "sucesso" ou "info" alternativo
COLOR_LIGHT_GRAY_BG = '#F0F0F0' # Um cinza bem leve para fundos que precisam ser muito claros, mas não branco puro.

# --- Funções Utilitárias ---
def validate_pdf(file_path):
    """Verifica se um arquivo PDF é válido e não está corrompido/vazio."""
    if not file_path or not os.path.exists(file_path):
        logging.warning(f"Caminho do PDF não existe: {file_path}")
        return False
    try:
        pdf_doc = fitz.open(file_path)
        if pdf_doc.page_count == 0:
            logging.warning(f"PDF sem páginas: {file_path}")
            pdf_doc.close()
            messagebox.showwarning(texts["error"], f"O arquivo {os.path.basename(file_path)} não contém páginas.")
            return False
        pdf_doc.close()
        logging.info(f"PDF validado com sucesso: {file_path}")
        return True
    except Exception as e:
        logging.error(f"PDF corrompido ou inválido: {file_path} - {str(e)}")
        messagebox.showerror(texts["error"], f"PDF corrompido ou formato não suportado:\n{os.path.basename(file_path)}\n({e})")
        return False

def select_file(title="Selecione o arquivo PDF"):
    """Abre uma caixa de diálogo para selecionar um único arquivo PDF."""
    file_path = filedialog.askopenfilename(title=title, filetypes=[("PDF files", "*.pdf")])
    if file_path and not validate_pdf(file_path):
        return None
    logging.debug(f"Arquivo selecionado: {file_path if file_path else 'Nenhum'}")
    return file_path

def select_files(title="Selecione os arquivos PDF"):
    """Abre uma caixa de diálogo para selecionar múltiplos arquivos PDF."""
    file_paths = filedialog.askopenfilenames(title=title, filetypes=[("PDF files", "*.pdf")])
    if not file_paths: return None
    
    valid_files = []
    invalid_files_basenames = []
    for fp in file_paths:
        if validate_pdf(fp):
            valid_files.append(fp)
        else:
            invalid_files_basenames.append(os.path.basename(fp))
    
    if invalid_files_basenames:
        messagebox.showwarning("Arquivos Inválidos",
                               f"Os seguintes arquivos foram ignorados (corrompidos ou inválidos):\n" +
                               f"\n - {', '.join(invalid_files_basenames)}")
    
    logging.debug(f"Arquivos selecionados válidos: {len(valid_files)}, inválidos: {len(invalid_files_basenames)}")
    return valid_files if valid_files else None

def get_file_size_mb(file_path):
    """Retorna o tamanho de um arquivo em MegaBytes."""
    try:
        return os.path.getsize(file_path) / (1024 * 1024)
    except OSError as e:
        logging.error(f"Erro ao obter tamanho do arquivo {file_path}: {e}")
        return 0

def validate_range(page_range_str, total_pages, allow_comma=True):
    """Valida uma string de intervalo de páginas (ex: '1-3', '5', '1,3,5')."""
    if not page_range_str.strip(): return False
    try:
        if allow_comma and ',' in page_range_str:
            pages = [int(p.strip()) for p in page_range_str.split(',')]
            return all(1 <= p <= total_pages for p in pages)
        elif '-' in page_range_str:
            parts = page_range_str.split('-')
            if len(parts) != 2: return False
            start, end = int(parts[0].strip()), int(parts[1].strip())
            return 1 <= start <= end <= total_pages
        else:
            page = int(page_range_str.strip())
            return 1 <= page <= total_pages
    except ValueError:
        return False

def parse_page_range(page_range_str, total_pages):
    """Converte uma string de intervalo de páginas em uma lista de índices baseados em zero."""
    indices = set()
    if not page_range_str.strip(): return []

    parts = page_range_str.replace(" ", "").split(',')
    for part in parts:
        if '-' in part:
            try:
                start_end_parts = part.split('-')
                if len(start_end_parts) == 2: 
                    start, end = map(int, start_end_parts)
                    if 1 <= start <= end <= total_pages:
                        indices.update(range(start - 1, end))
                else: 
                    logging.warning(f"Intervalo malformado '{part}' em '{page_range_str}'")
            except ValueError:
                logging.warning(f"Intervalo inválido (não numérico) '{part}' em '{page_range_str}'")
                continue
        else:
            try:
                page = int(part)
                if 1 <= page <= total_pages:
                    indices.add(page - 1)
            except ValueError:
                logging.warning(f"Número de página inválido '{part}' em '{page_range_str}'")
                continue
    return sorted(list(indices))

def validate_parts(parts_str, total_pages):
    """Valida o número de partes para a divisão de PDF."""
    try:
        num_parts = int(parts_str)
        return 2 <= num_parts <= total_pages
    except ValueError:
        return False

def create_tooltip(widget, text):
    """Cria tooltips customizados para os widgets."""
    tooltip_active = False
    tooltip_window = None
    
    def enter(event):
        nonlocal tooltip_active, tooltip_window
        if tooltip_active: return
        tooltip_active = True
        x, y, _, _ = widget.bbox("insert")
        x += widget.winfo_rootx() + 25
        y += widget.winfo_rooty() + 25
        tooltip_window = tk.Toplevel(widget)
        tooltip_window.wm_overrideredirect(True)
        tooltip_window.wm_geometry(f"+{x}+{y}")
        try:
            tooltip_window.attributes("-alpha", 0.0) # Inicia transparente
        except tk.TclError: pass
        # Fundo do tooltip em Noite Estrelada e texto em Luz Pura da Lua
        label = ttkb.Label(tooltip_window, text=text, background=COLOR_NIGHT_SKY, foreground=COLOR_MOON_LIGHT, relief="solid", borderwidth=1, font=("Verdana", 9, "italic"), padding=(6,4), wraplength=280)
        label.pack()
        try:
            tooltip_window.attributes("-alpha", 0.95) # Fade in
        except tk.TclError: pass
        widget.after_id = widget.after(4000, lambda: fade_out(True)) # Fecha após 4 segundos

    def fade_out(triggered_by_timer=False):
        nonlocal tooltip_active, tooltip_window
        if triggered_by_timer: # Se foi ativado pelo timer, cancela o bind de Leave (se ele existir)
             pass
        elif hasattr(widget, 'after_id') and widget.after_id: # Se o mouse sair, cancela o timer
            widget.after_cancel(widget.after_id)
            widget.after_id = None

        if tooltip_active and tooltip_window:
            try: tooltip_window.destroy()
            except tk.TclError: pass # Ignora erro se a janela já foi destruída
            tooltip_active = False
            tooltip_window = None
            
    def leave(event):
        fade_out(False)

    widget.bind("<Enter>", enter)
    widget.bind("<Leave>", leave)

def get_description_for_function(function_key):
    """Retorna uma descrição detalhada para a função baseada na chave."""
    # Mapeamento direto para descrições específicas de prévia
    desc_key_map = {
        "exclude_pages": "interactive_page_range_description_exclude",
        "select_pages": "interactive_page_range_description_select",
        "add_pages": "interactive_add_description",
        "replace_pages": "interactive_replace_description",
        "split_pdf": "tooltip_split", # Usando tooltip como fallback para descrição curta
        "merge_pdf": "tooltip_merge",
        "compress_pdf": "tooltip_compress",
        "anonymize_pdf": "tooltip_anonymize",
        "convert_pdf": "tooltip_convert",
        "pdf_jpg": "tooltip_pdf_jpg"
    }
    desc_key = desc_key_map.get(function_key)
    
    # Tenta buscar a descrição mapeada
    if desc_key and desc_key in texts:
        return texts[desc_key]
    
    # Fallback: tenta buscar no help_text se não houver mapeamento direto
    help_content = texts.get("help_text", "")
    lines = help_content.split('\n')
    search_term_display_text = texts.get(function_key, function_key) # Usa o texto do botão como termo de busca
    for line in lines:
        if line.strip().startswith(f"- {search_term_display_text}"):
            parts = line.split(':', 1)
            if len(parts) > 1:
                return parts[1].strip()
    return "Defina os parâmetros para a operação." # Descrição genérica se nada for encontrado

def clear_options_frame():
    """Limpa todos os widgets do frame de opções."""
    for widget in options_frame.winfo_children():
        widget.destroy()
    logging.debug("Frame de opções limpo.")

# --- Funções de Prévia Interativa Base ---
def render_pdf_page_to_image(page_obj, max_width=MAX_PREVIEW_IMG_WIDTH_SINGLE):
    """Renderiza uma página PDF para uma imagem Tkinter PhotoImage."""
    global preview_images
    try:
        # Usa um fator de resolução para melhor qualidade (ajustável)
        matrix = fitz.Matrix(PREVIEW_RESOLUTION_FACTOR, PREVIEW_RESOLUTION_FACTOR)
        pix = page_obj.get_pixmap(matrix=matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        if pix.width == 0: return None # Evitar divisão por zero
        
        # Redimensiona para a largura máxima desejada, mantendo a proporção
        ratio = max_width / pix.width
        new_height = int(pix.height * ratio)
        img_resized = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
        
        photo = ImageTk.PhotoImage(img_resized)
        preview_images.append(photo) # Manter referência para evitar garbage collection
        return photo
    except Exception as e:
        logging.error(f"Erro ao renderizar página {page_obj.number + 1} para prévia: {e}")
        return None

def create_interactive_preview_window_base(title_key, num_pdf_panels=1):
    """Cria a janela base para prévias interativas."""
    global preview_images
    preview_images = [] # Limpa a lista de referências de imagens para a nova janela

    preview_window = tk.Toplevel(root)
    preview_window.title(texts[title_key])
    geo_width = 900 if num_pdf_panels == 1 else 1200 # Ajusta largura da janela
    preview_window.geometry(f"{geo_width}x700")
    preview_window.minsize(700 if num_pdf_panels==1 else 1000, 600)
    
    # Fundo da janela de prévia: Noite Estrelada
    preview_window.configure(bg=COLOR_NIGHT_SKY) 

    # Frames para organização da janela (todos com fundo Noite Estrelada)
    top_input_frame = ttkb.Frame(preview_window, padding=(10, 10), style='PreviewDark.TFrame') 
    top_input_frame.pack(side="top", fill="x", padx=5, pady=5)

    main_preview_area_frame = ttkb.Frame(preview_window, style='PreviewDark.TFrame') 
    main_preview_area_frame.pack(side="top", fill="both", expand=True, padx=5, pady=(0,5))

    bottom_action_frame = ttkb.Frame(preview_window, padding=(10,10), style='PreviewDark.TFrame') 
    bottom_action_frame.pack(side="bottom", fill="x", padx=5, pady=5)

    return preview_window, top_input_frame, main_preview_area_frame, bottom_action_frame

def setup_scrollable_canvas_in_frame(parent_frame, label_text="Prévia PDF"):
    """Configura um canvas rolável dentro de um frame para exibir prévias de páginas."""
    # LabelFrame com fundo Noite Estrelada, texto e borda Terracota
    preview_panel = ttkb.LabelFrame(parent_frame, text=label_text, style='PreviewInfo.TLabelframe', padding=(5,5))
    preview_panel.pack(side="left", fill="both", expand=True, padx=(0,5) if len(parent_frame.winfo_children()) > 0 else (0,0))

    # Container do canvas com fundo Noite Estrelada
    canvas_container = ttkb.Frame(preview_panel, style='PreviewDark.TFrame', borderwidth=1, relief="sunken")
    canvas_container.pack(fill="both", expand=True)

    # Fundo do canvas: Luz Pura da Lua (#ffffff) para visibilidade do PDF
    canvas = tk.Canvas(canvas_container, bg=COLOR_MOON_LIGHT, highlightthickness=0)
    # Scrollbar: Terracota do Cerrado para a cor do polegar
    scrollbar = ttkb.Scrollbar(canvas_container, orient="vertical", command=canvas.yview, bootstyle="round-info")
    # Configura o estilo da scrollbar diretamente
    style.configure("round-info", troughcolor=COLOR_NIGHT_SKY, background=COLOR_TERRACOTTA) 
    
    # Frame interno do canvas com fundo Luz Pura da Lua
    scrollable_content_frame = ttkb.Frame(canvas, background=COLOR_MOON_LIGHT) 

    # Configura o scrollregion do canvas quando o conteúdo é redimensionado
    scrollable_content_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
    canvas.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)
    canvas.create_window((0, 0), window=scrollable_content_frame, anchor="nw", tags="frame_content")

    # Adiciona suporte para rolagem com a roda do mouse
    def _on_mousewheel(event, c=canvas):
        if not c.winfo_exists(): 
            return
        delta = 0
        if event.num == 4: 
            delta = -1 
        elif event.num == 5: 
            delta = 1 
        else: 
            delta = -1 if event.delta > 0 else 1
            if abs(event.delta) < 100 and abs(event.delta) > 2: 
                delta = int(event.delta /10) if abs(event.delta)>5 else event.delta
            elif abs(event.delta) <=2: 
                delta = 0
        if delta != 0:
            c.yview_scroll(delta , "units")

    # Bind da roda do mouse para Windows/macOS e Linux
    for widget_to_bind in [canvas, scrollable_content_frame]: 
        widget_to_bind.bind("<MouseWheel>", lambda e, c_arg=canvas: _on_mousewheel(e, c_arg), add="+") 
        widget_to_bind.bind("<Button-4>", lambda e, c_arg=canvas: _on_mousewheel(type('event', (object,), {'num': 4, 'delta': 0}), c_arg), add="+")
        widget_to_bind.bind("<Button-5>", lambda e, c_arg=canvas: _on_mousewheel(type('event', (object,), {'num': 5, 'delta': 0}), c_arg), add="+")
        
    return canvas, scrollable_content_frame

# --- Funções de Prévia Interativa Específicas ---

def open_interactive_preview_single_pdf(file_path, title_key, entry_prompt_key, entry_desc_key, action_callback):
    """Abre a janela de prévia interativa para operações com um único PDF (Excluir, Selecionar)."""
    global preview_images
    preview_window, top_input_frame, main_preview_area_frame, bottom_action_frame = \
        create_interactive_preview_window_base(title_key, num_pdf_panels=1)

    # --- Painel de Entrada ---
    # Labels com fundo Noite Estrelada e texto Luz Pura da Lua
    input_main_label = ttkb.Label(top_input_frame, text=texts[entry_prompt_key], background=COLOR_NIGHT_SKY, foreground=COLOR_MOON_LIGHT, font=("Verdana", 10))
    input_main_label.pack(side="left", padx=(0,5), pady=5, anchor="w")
    
    # Entry com fundo Luz Pura da Lua e texto Noite Estrelada
    page_range_entry = ttkb.Entry(top_input_frame, width=25, font=("Verdana", 10), fieldbackground=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY, bordercolor=COLOR_TERRACOTTA)
    page_range_entry.pack(side="left", padx=5, pady=5, anchor="w")
    
    # Descrição com fundo Noite Estrelada e texto Terracota
    description_text = get_description_for_function(entry_desc_key) 
    description_label = ttkb.Label(top_input_frame, text=description_text, background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA, wraplength=350, justify="left", font=("Verdana", 8))
    description_label.pack(side="left", padx=5, pady=5, fill="x", expand=True, anchor="w")

    # --- Painel de Prévia ---
    canvas, scroll_content = setup_scrollable_canvas_in_frame(main_preview_area_frame, texts["preview_info_original_pdf"].format(filename=os.path.basename(file_path)))

    # Usar um dicionário para a referência do documento PDF para que possa ser atualizado em funções aninhadas
    pdf_doc_ref = {'doc': None, 'total_pages': 0} 
    try:
        temp_doc = fitz.open(file_path)
        pdf_doc_ref['total_pages'] = temp_doc.page_count
        temp_doc.close() 
    except Exception as e: 
        logging.error(f"Erro crítico ao abrir PDF {file_path} inicialmente para prévia: {e}")
        messagebox.showerror(texts["error"], f"Não foi possível abrir o PDF: {os.path.basename(file_path)}\nVerifique se o arquivo não está corrompido.")
        preview_window.destroy()
        return
    
    if pdf_doc_ref['total_pages'] == 0:
        messagebox.showerror(texts["error"], f"O PDF {os.path.basename(file_path)} não contém páginas.")
        preview_window.destroy()
        return

    def _update_preview_display(event=None):
        global preview_images
        preview_images = [] 
        for widget in scroll_content.winfo_children(): 
            widget.destroy()

        range_str = page_range_entry.get()
        
        # Reabre o documento se ele estiver fechado (para garantir acesso consistente)
        if not pdf_doc_ref['doc'] or pdf_doc_ref['doc'].is_closed:
            try:
                pdf_doc_ref['doc'] = fitz.open(file_path)
                pdf_doc_ref['total_pages'] = pdf_doc_ref['doc'].page_count 
            except Exception as e:
                logging.error(f"Erro ao reabrir PDF {file_path} para prévia: {e}")
                ttkb.Label(scroll_content, text="Erro ao carregar PDF para prévia.", background=COLOR_MOON_LIGHT, foreground="red").pack(pady=5) 
                return
        
        pdf_doc = pdf_doc_ref['doc']
        total_pages = pdf_doc_ref['total_pages']

        if not validate_range(range_str, total_pages, allow_comma=True):
            ttkb.Label(scroll_content, text=texts["invalid_range"].format(total=total_pages), background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5) 
            canvas.yview_moveto(0) 
            scroll_content.update_idletasks() 
            canvas.config(scrollregion=canvas.bbox("all")) 
            return

        indices_to_display = parse_page_range(range_str, total_pages)
        
        # Texto descritivo para a operação
        operation_type_text_map = {
            "preview_title_exclude": texts.get("interactive_page_range_description_exclude_label", "Páginas a serem EXCLUÍDAS:"),
            "preview_title_select": texts.get("interactive_page_range_description_select_label", "Páginas a serem SELECIONADAS:")
        }
        operation_type_text = operation_type_text_map.get(title_key, "Páginas Afetadas:")
        
        if operation_type_text:
            ttkb.Label(scroll_content, text=operation_type_text, background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(5,0), anchor="w", padx=5) 

        if not indices_to_display:
             ttkb.Label(scroll_content, text="Nenhuma página corresponde ao intervalo digitado.", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5) 
        
        # Exibe as imagens de prévia
        for i in indices_to_display[:MAX_PREVIEW_IMAGES_SHOWN]:
            if 0 <= i < total_pages: 
                try:
                    page = pdf_doc.load_page(i)
                    photo = render_pdf_page_to_image(page, max_width=MAX_PREVIEW_IMG_WIDTH_SINGLE)
                    if photo:
                        ttkb.Label(scroll_content, image=photo, background=COLOR_MOON_LIGHT).pack(pady=10, padx=10)
                        ttkb.Label(scroll_content, text=texts["preview_page_label"].format(num=i + 1), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(0, 10))
                except Exception as render_err: 
                    logging.error(f"Erro ao renderizar página {i+1} para prévia (single_pdf): {render_err}")
                    ttkb.Label(scroll_content, text=f"Erro ao renderizar pág. {i+1}", background=COLOR_MOON_LIGHT, foreground="red").pack(pady=5)

        if len(indices_to_display) > MAX_PREVIEW_IMAGES_SHOWN:
             ttkb.Label(scroll_content, text=f"... e mais {len(indices_to_display) - MAX_PREVIEW_IMAGES_SHOWN} página(s) (não exibidas).", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5) 
        
        canvas.yview_moveto(0) 
        scroll_content.update_idletasks() 
        canvas.config(scrollregion=canvas.bbox("all")) 
        
    def on_confirm_action():
        """Função chamada ao confirmar a operação na janela de prévia."""
        current_total_pages = 0
        if not pdf_doc_ref['doc'] or pdf_doc_ref['doc'].is_closed:
            try:
                temp_doc_confirm = fitz.open(file_path)
                current_total_pages = temp_doc_confirm.page_count
                temp_doc_confirm.close()
            except Exception as e_confirm:
                logging.error(f"Erro ao abrir PDF {file_path} na confirmação: {e_confirm}")
                messagebox.showerror(texts["error"], "Erro ao acessar o PDF para confirmar a operação.")
                preview_window.destroy() 
                return
        else: 
            current_total_pages = pdf_doc_ref['doc'].page_count
            pdf_doc_ref['doc'].close() 
            pdf_doc_ref['doc'] = None

        action_callback(file_path, page_range_entry.get(), current_total_pages)
        preview_window.destroy() 

    def on_close_preview_window():
        """Fecha o documento PDF aberto na prévia ao fechar a janela."""
        if pdf_doc_ref['doc'] and not pdf_doc_ref['doc'].is_closed:
            pdf_doc_ref['doc'].close()
            logging.debug(f"PDF {file_path} fechado ao sair da prévia (single_pdf).")
        preview_window.destroy()

    # Binds para atualizar a prévia
    page_range_entry.bind("<FocusOut>", _update_preview_display)
    page_range_entry.bind("<Return>", _update_preview_display)
    
    # Botões de ação
    confirm_btn = ttkb.Button(bottom_action_frame, text=texts["preview_confirm"], background=COLOR_TERRACOTTA, foreground=COLOR_NIGHT_SKY, font=("Verdana", 10, "bold"))
    confirm_btn.pack(side="left", padx=10, pady=5, fill="x", expand=True)
    cancel_btn = ttkb.Button(bottom_action_frame, text=texts["preview_cancel"], background=COLOR_DARK_EARTH, foreground=COLOR_MOON_LIGHT, font=("Verdana", 10, "bold"))
    cancel_btn.pack(side="right", padx=10, pady=5, fill="x", expand=True)
    
    preview_window.protocol("WM_DELETE_WINDOW", on_close_preview_window) 
    page_range_entry.focus_set() 
    _update_preview_display() 

def open_interactive_preview_dual_pdf(file_path_orig, file_path_second, title_key, entry_prompt_key_main, entry_desc_key, action_callback):
    """Abre a janela de prévia interativa para operações com dois PDFs (Acrescentar, Substituir)."""
    global preview_images
    preview_window, top_input_frame, main_preview_area_frame, bottom_action_frame = \
        create_interactive_preview_window_base(title_key, num_pdf_panels=2)

    # Dicionários para referências dos documentos PDF
    pdf_doc_orig_ref = {'doc': None, 'total_pages': 0}
    pdf_doc_second_ref = {'doc': None, 'total_pages': 0}

    try:
        pdf_doc_orig_ref['doc'] = fitz.open(file_path_orig)
        pdf_doc_orig_ref['total_pages'] = pdf_doc_orig_ref['doc'].page_count
        pdf_doc_second_ref['doc'] = fitz.open(file_path_second)
        pdf_doc_second_ref['total_pages'] = pdf_doc_second_ref['doc'].page_count
    except Exception as e:
        logging.error(f"Erro crítico ao abrir PDFs para prévia dual: {e}")
        messagebox.showerror(texts["error"], f"Não foi possível abrir um dos PDFs:\n{os.path.basename(file_path_orig)} ou {os.path.basename(file_path_second)}\nVerifique se não estão corrompidos.")
        preview_window.destroy()
        return

    # Verificações adicionais para casos de PDFs vazios
    if title_key == "preview_title_add" and pdf_doc_orig_ref['total_pages'] == 0 and pdf_doc_second_ref['total_pages'] == 0:
        messagebox.showerror(texts["error"], f"Pelo menos um dos PDFs precisa ter páginas para a operação de acrescentar.")
        preview_window.destroy()
        return
    elif title_key == "preview_title_replace" and (pdf_doc_orig_ref['total_pages'] == 0 or pdf_doc_second_ref['total_pages'] == 0):
        messagebox.showerror(texts["error"], "Ambos os PDFs (original e de substituição) precisam ter páginas para a operação de substituir.")
        preview_window.destroy()
        return

    # --- Painel de Entrada ---
    input_group_main = ttkb.Frame(top_input_frame, style='PreviewDark.TFrame') 
    input_group_main.pack(side="left", fill="x", expand=True, padx=(0,10))

    ttkb.Label(input_group_main, text=texts["interactive_file_original_label"], background=COLOR_NIGHT_SKY, foreground=COLOR_MOON_LIGHT).pack(anchor="w")
    ttkb.Label(input_group_main, text=os.path.basename(file_path_orig), background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA, font=("Verdana", 9, "bold")).pack(anchor="w", pady=(0,2))
    
    prompt_text_main = texts[entry_prompt_key_main]
    if "{total_original}" in prompt_text_main: 
        prompt_text_main = prompt_text_main.format(total_original=pdf_doc_orig_ref['total_pages'])
    
    ttkb.Label(input_group_main, text=prompt_text_main, background=COLOR_NIGHT_SKY, foreground=COLOR_MOON_LIGHT).pack(anchor="w", pady=(5,0))
    param_entry = ttkb.Entry(input_group_main, width=20, font=("Verdana", 10), fieldbackground=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY, bordercolor=COLOR_TERRACOTTA)
    param_entry.pack(anchor="w", pady=(2,5))

    input_group_second = ttkb.Frame(top_input_frame, style='PreviewDark.TFrame') 
    input_group_second.pack(side="left", fill="x", expand=True, padx=(10,0))
    
    ttkb.Label(input_group_second, text=texts["interactive_file_second_label"], background=COLOR_NIGHT_SKY, foreground=COLOR_MOON_LIGHT).pack(anchor="w")
    ttkb.Label(input_group_second, text=os.path.basename(file_path_second), background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA, font=("Verdana", 9, "bold")).pack(anchor="w", pady=(0,2))
    ttkb.Label(input_group_second, text=get_description_for_function(entry_desc_key), background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA, wraplength=300, justify="left", font=("Verdana", 8)).pack(anchor="w", pady=(5,0))

    # --- Painéis de Prévia ---
    canvas_orig, scroll_content_orig = setup_scrollable_canvas_in_frame(main_preview_area_frame, texts["preview_info_original_pdf"].format(filename=os.path.basename(file_path_orig)))
    canvas_second, scroll_content_second = setup_scrollable_canvas_in_frame(main_preview_area_frame, texts["preview_info_second_pdf"].format(filename=os.path.basename(file_path_second)))

    def _update_dual_preview_display(event=None):
        global preview_images
        preview_images = [] 
        for widget in scroll_content_orig.winfo_children(): widget.destroy()
        for widget in scroll_content_second.winfo_children(): widget.destroy()

        param_str = param_entry.get()
        
        # Reabre os documentos se estiverem fechados
        if pdf_doc_orig_ref['doc'].is_closed: pdf_doc_orig_ref['doc'] = fitz.open(file_path_orig)
        if pdf_doc_second_ref['doc'].is_closed: pdf_doc_second_ref['doc'] = fitz.open(file_path_second)
        pdf_orig = pdf_doc_orig_ref['doc']
        pdf_second = pdf_doc_second_ref['doc']
        total_orig = pdf_doc_orig_ref['total_pages']
        total_sec = pdf_doc_second_ref['total_pages']

        if title_key == "preview_title_add":
            try:
                insert_after_page_idx_user = int(param_str) 
                if not (0 <= insert_after_page_idx_user <= total_orig):
                    ttkb.Label(scroll_content_orig, text=texts["invalid_insert_point"].format(total=total_orig), background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)
                    return
            except ValueError:
                ttkb.Label(scroll_content_orig, text="Ponto de inserção inválido.", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)
                return
            
            # Exibe a página antes do ponto de inserção
            pages_to_show_orig_context = []
            if insert_after_page_idx_user > 0 and insert_after_page_idx_user <= total_orig :
                pages_to_show_orig_context.append(insert_after_page_idx_user - 1)
            
            marker_shown = False
            for idx in pages_to_show_orig_context[:MAX_PREVIEW_IMAGES_SHOWN]: 
                if 0 <= idx < total_orig:
                    page = pdf_orig.load_page(idx)
                    photo = render_pdf_page_to_image(page, MAX_PREVIEW_IMG_WIDTH_DUAL)
                    if photo:
                        ttkb.Label(scroll_content_orig, image=photo, background=COLOR_MOON_LIGHT).pack(pady=5)
                        ttkb.Label(scroll_content_orig, text=texts["preview_page_of_label"].format(num=idx + 1, pdf_name="Orig."), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack()
            
            # Marcador de inserção
            ttkb.Label(scroll_content_orig, text=texts["preview_add_insert_label"], background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA, font="Verdana 10 bold").pack(pady=10, fill="x")
            marker_shown = True
            
            # Exibe a página depois do ponto de inserção
            idx_after = insert_after_page_idx_user 
            if idx_after < total_orig:
                    page = pdf_orig.load_page(idx_after)
                    photo = render_pdf_page_to_image(page, MAX_PREVIEW_IMG_WIDTH_DUAL)
                    if photo:
                        ttkb.Label(scroll_content_orig, image=photo, background=COLOR_MOON_LIGHT).pack(pady=5)
                        ttkb.Label(scroll_content_orig, text=texts["preview_page_of_label"].format(num=idx_after + 1, pdf_name="Orig."), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack()
            elif not marker_shown and total_orig > 0: 
                 ttkb.Label(scroll_content_orig, text="(Inserir no final do PDF Original)", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)
            elif total_orig == 0 : 
                 ttkb.Label(scroll_content_orig, text="(PDF Original está vazio)", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)


            # Prévia do segundo PDF (a ser inserido)
            if total_sec == 0:
                ttkb.Label(scroll_content_second, text="(PDF Adicional está vazio)", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)
            for i in range(min(total_sec, MAX_PREVIEW_IMAGES_SHOWN)):
                page = pdf_second.load_page(i)
                photo = render_pdf_page_to_image(page, MAX_PREVIEW_IMG_WIDTH_DUAL)
                if photo:
                    ttkb.Label(scroll_content_second, image=photo, background=COLOR_MOON_LIGHT).pack(pady=5)
                    ttkb.Label(scroll_content_second, text=texts["preview_page_of_label"].format(num=i + 1, pdf_name="Adic."), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack()
            if total_sec > MAX_PREVIEW_IMAGES_SHOWN:
                ttkb.Label(scroll_content_second, text=f"... e mais {total_sec - MAX_PREVIEW_IMAGES_SHOWN} pág.", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)

        elif title_key == "preview_title_replace":
            if not validate_range(param_str, total_orig, allow_comma=False): 
                ttkb.Label(scroll_content_orig, text=texts["invalid_range"].format(total=total_orig), background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)
                return
            
            indices_to_replace_orig = parse_page_range(param_str, total_orig)
            if not indices_to_replace_orig:
                 ttkb.Label(scroll_content_orig, text="Nenhuma página válida no intervalo.", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)
                 return

            # Exibe as páginas do PDF original a serem substituídas e as páginas do PDF substituto
            for i, orig_idx in enumerate(indices_to_replace_orig[:MAX_PREVIEW_IMAGES_SHOWN]):
                if 0 <= orig_idx < total_orig:
                    page_o = pdf_orig.load_page(orig_idx)
                    photo_o = render_pdf_page_to_image(page_o, MAX_PREVIEW_IMG_WIDTH_DUAL)
                    if photo_o:
                        ttkb.Label(scroll_content_orig, image=photo_o, background=COLOR_MOON_LIGHT).pack(pady=5)
                        ttkb.Label(scroll_content_orig, text=texts["preview_replace_original_label"].format(num=orig_idx + 1), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack()
                
                if i < total_sec: 
                    page_r = pdf_second.load_page(i)
                    photo_r = render_pdf_page_to_image(page_r, MAX_PREVIEW_IMG_WIDTH_DUAL)
                    if photo_r:
                        ttkb.Label(scroll_content_second, image=photo_r, background=COLOR_MOON_LIGHT).pack(pady=5)
                        ttkb.Label(scroll_content_second, text=texts["preview_replace_new_label"].format(num=i + 1), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack()
                else: 
                    ttkb.Label(scroll_content_second, text=f"(Sem pág. {i+1} para substituir)", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)
                
                if i < min(len(indices_to_replace_orig), MAX_PREVIEW_IMAGES_SHOWN) -1 : 
                     ttkb.Separator(scroll_content_orig, orient=HORIZONTAL, background=COLOR_TERRACOTTA).pack(pady=5, fill="x", padx=10)
                     ttkb.Separator(scroll_content_second, orient=HORIZONTAL, background=COLOR_TERRACOTTA).pack(pady=5, fill="x", padx=10)


            if len(indices_to_replace_orig) > MAX_PREVIEW_IMAGES_SHOWN:
                ttkb.Label(scroll_content_orig, text=f"... e mais {len(indices_to_replace_orig) - MAX_PREVIEW_IMAGES_SHOWN} pág.", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH).pack(pady=5)
            if len(indices_to_replace_orig) > total_sec and total_sec > 0: 
                 ttkb.Label(scroll_content_second, text=f"Aviso: {len(indices_to_replace_orig)} pág. para substituir, mas PDF substituto tem apenas {total_sec} pág.", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)
            elif total_sec == 0:
                 ttkb.Label(scroll_content_second, text="(PDF de substituição está vazio)", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA).pack(pady=5)

        # Ajusta as scrollregions dos canvases
        for canvas_item in [canvas_orig, canvas_second]:
            canvas_item.yview_moveto(0)
            frame_tag = canvas_item.find_withtag("frame_content")
            if frame_tag:
                content_f = canvas_item.nametowidget(canvas_item.itemcget(frame_tag[0], "window"))
                content_f.update_idletasks()
                canvas_item.config(scrollregion=canvas_item.bbox("all"))

    def on_confirm_action_dual():
        """Função chamada ao confirmar a operação com dois PDFs."""
        on_close_preview_dual() 
        action_callback(file_path_orig, file_path_second, param_entry.get(), 
                        pdf_doc_orig_ref['total_pages'], pdf_doc_second_ref['total_pages'])

    def on_close_preview_dual():
        """Fecha os documentos PDF abertos na prévia ao fechar a janela."""
        if pdf_doc_orig_ref['doc'] and not pdf_doc_orig_ref['doc'].is_closed:
            pdf_doc_orig_ref['doc'].close()
        if pdf_doc_second_ref['doc'] and not pdf_doc_second_ref['doc'].is_closed:
            pdf_doc_second_ref['doc'].close()
        logging.debug(f"PDFs {file_path_orig} e {file_path_second} fechados ao sair da prévia (dual).")
        preview_window.destroy()

    # Binds para atualizar a prévia
    param_entry.bind("<FocusOut>", _update_dual_preview_display)
    param_entry.bind("<Return>", _update_dual_preview_display)

    # Botões de ação
    confirm_btn = ttkb.Button(bottom_action_frame, text=texts["preview_confirm"], background=COLOR_TERRACOTTA, foreground=COLOR_NIGHT_SKY, font=("Verdana", 10, "bold"))
    confirm_btn.pack(side="left", padx=10, pady=5, fill="x", expand=True)
    cancel_btn = ttkb.Button(bottom_action_frame, text=texts["preview_cancel"], background=COLOR_DARK_EARTH, foreground=COLOR_MOON_LIGHT, font=("Verdana", 10, "bold"))
    cancel_btn.pack(side="right", padx=10, pady=5, fill="x", expand=True)
    
    preview_window.protocol("WM_DELETE_WINDOW", on_close_preview_dual) 
    param_entry.focus_set()
    _update_dual_preview_display()

    # --- Funções de Botão da UI Principal MODIFICADAS ---
def on_exclude_button():
    """Lógica para o botão 'Excluir Página(s)', inicia a prévia interativa."""
    file_path = select_file(title="Selecione o PDF para Excluir Páginas")
    if not file_path: return
    clear_options_frame() # Limpa o painel de opções antes de abrir a prévia
    open_interactive_preview_single_pdf(file_path, 
                                         "preview_title_exclude", 
                                         "interactive_exclude_prompt", 
                                         "exclude_pages", 
                                         exclude_pages_action)

def select_pages(): 
    """Lógica para o botão 'Selecionar Página(s)', inicia a prévia interativa."""
    file_path = select_file(title="Selecione o PDF para Selecionar Páginas")
    if not file_path: return
    clear_options_frame()
    open_interactive_preview_single_pdf(file_path, 
                                         "preview_title_select", 
                                         "interactive_select_prompt", 
                                         "select_pages", 
                                         select_pages_action)

def add_selected_pages(): 
    """Lógica para o botão 'Acrescentar Página(s)', inicia a prévia interativa com dois PDFs."""
    clear_options_frame()
    file_path_orig = select_file(title=texts["interactive_file_original_label"])
    if not file_path_orig: return

    file_path_second = select_file(title=texts["interactive_file_second_label"])
    if not file_path_second: return
    
    # Prevenção para não usar o mesmo arquivo duas vezes na adição
    if file_path_orig == file_path_second:
        messagebox.showwarning("Arquivos Iguais", "O PDF original e o PDF a ser acrescentado não podem ser o mesmo arquivo para esta operação.")
        return

    open_interactive_preview_dual_pdf(file_path_orig, file_path_second,
                                      "preview_title_add", 
                                      "interactive_add_prompt", 
                                      "add_pages", 
                                      add_pages_action)

def on_replace_button(): 
    """Lógica para o botão 'Substituir Página(s)', inicia a prévia interativa com dois PDFs."""
    clear_options_frame()
    file_path_orig = select_file(title=texts["interactive_file_original_label"])
    if not file_path_orig: return

    file_path_second = select_file(title=texts["interactive_file_second_label"])
    if not file_path_second: return

    # Prevenção para não usar o mesmo arquivo duas vezes na substituição
    if file_path_orig == file_path_second:
        messagebox.showwarning("Arquivos Iguais", "O PDF original e o PDF de substituição não podem ser o mesmo arquivo para esta operação.")
        return

    open_interactive_preview_dual_pdf(file_path_orig, file_path_second,
                                      "preview_title_replace", 
                                      "interactive_replace_prompt", 
                                      "replace_pages", 
                                      replace_pages_action)

# --- Funções de AÇÃO (Manipulação de PDF) ---
def exclude_pages_action(file_path, page_range_str, total_pages_original_from_preview):
    """Executa a exclusão de páginas de um PDF."""
    logging.info(f"Iniciando exclusão de páginas para {file_path}, range: '{page_range_str}'")
    
    indices_to_exclude = parse_page_range(page_range_str, total_pages_original_from_preview)
    if not indices_to_exclude and page_range_str.strip(): # Se o range não é vazio, mas não gerou índices válidos
        messagebox.showerror(texts["error"], texts["invalid_range"].format(total=total_pages_original_from_preview) + "\nNenhuma página válida no intervalo para excluir.")
        return
    if not indices_to_exclude: # Se o range é vazio ou não tem páginas válidas
        messagebox.showinfo("Nenhuma Ação", "Nenhum intervalo de páginas válido fornecido para exclusão.")
        return

    input_pdf_stream = None
    try:
        input_pdf_stream = open(file_path, "rb")
        input_pdf = PdfReader(input_pdf_stream)
        actual_total_pages = len(input_pdf.pages)
        
        # Re-validação da contagem de páginas para o caso do arquivo ter mudado
        if abs(actual_total_pages - total_pages_original_from_preview) > 1 : # Tolerância de 1 para evitar falsos positivos
             logging.warning(f"Contagem de páginas divergente para {file_path}. Prévia: {total_pages_original_from_preview}, Atual: {actual_total_pages}")
             if not messagebox.askyesno("Aviso de Alteração", f"O número de páginas em {os.path.basename(file_path)} mudou de {total_pages_original_from_preview} para {actual_total_pages} desde a prévia.\nDeseja continuar a exclusão com o intervalo '{page_range_str}' aplicado ao novo total de páginas?"):
                 return
             # Re-parse do range com o novo total de páginas
             indices_to_exclude = parse_page_range(page_range_str, actual_total_pages)
             if not indices_to_exclude:
                 messagebox.showerror(texts["error"], "O intervalo fornecido não é mais válido com o novo total de páginas.")
                 return

        output_pdf = PdfWriter()
        progress_var.set(0)
        progress_bar["maximum"] = actual_total_pages
        start_time = time.time()
        included_page_count = 0

        for i in range(actual_total_pages):
            elapsed = time.time() - start_time
            eta = (elapsed / (i + 1)) * (actual_total_pages - (i + 1)) if i > 0 else 0
            status_label.config(text=texts["processing_page"].format(current=i + 1, total=actual_total_pages, eta=eta))
            
            if i not in indices_to_exclude:
                try:
                    output_pdf.add_page(input_pdf.pages[i])
                    included_page_count += 1
                except Exception as page_err: # Erro mais geral para renderização de página
                    logging.error(f"Erro ao adicionar página {i+1} de {file_path} durante exclusão: {page_err}")
            
            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()

        if included_page_count == 0:
             messagebox.showwarning("Resultado Vazio", "Todas as páginas foram selecionadas para exclusão. Nenhum arquivo foi gerado.")
             status_label.config(text="")
             progress_var.set(0)
             return

        output_filename = filedialog.asksaveasfilename(
            title="Salvar PDF com páginas excluídas como...",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"{os.path.splitext(os.path.basename(file_path))[0]}_excluido.pdf")
        
        if not output_filename:
            status_label.config(text="Operação cancelada pelo usuário.")
            progress_var.set(0)
            return

        with open(output_filename, "wb") as out_file:
            output_pdf.write(out_file)

        status_label.config(text="")
        logging.info(f"Páginas excluídas com sucesso. Arquivo salvo como {output_filename}")
        messagebox.showinfo(texts["success"], f"Páginas excluídas com sucesso!\nArquivo salvo em: {output_filename}")

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Arquivo não encontrado em exclude_pages_action: {file_path}")
        messagebox.showerror(texts["error"], f"Arquivo não encontrado: {os.path.basename(file_path)}")
    except Exception as e:
        status_label.config(text="Erro na operação.")
        logging.error(f"Erro em exclude_pages_action para {file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao excluir páginas: {str(e)}\nConsulte o log para mais detalhes.")
    finally:
        progress_var.set(0)
        if input_pdf_stream and not input_pdf_stream.closed: # Garante que o stream seja fechado
            input_pdf_stream.close()


def select_pages_action(file_path, page_range_str, total_pages_original_from_preview):
    """Executa a seleção (extração) de páginas de um PDF."""
    logging.info(f"Iniciando seleção de páginas para {file_path}, range: '{page_range_str}'")
    
    indices_to_select = parse_page_range(page_range_str, total_pages_original_from_preview)
    if not indices_to_select and page_range_str.strip():
        messagebox.showerror(texts["error"], texts["invalid_range"].format(total=total_pages_original_from_preview) + "\nNenhuma página válida no intervalo para selecionar.")
        return
    if not indices_to_select:
        messagebox.showinfo("Nenhuma Ação", "Nenhum intervalo de páginas válido fornecido para seleção.")
        return
    
    input_pdf_stream = None
    try:
        input_pdf_stream = open(file_path, "rb")
        input_pdf = PdfReader(input_pdf_stream)
        actual_total_pages = len(input_pdf.pages)
        
        if abs(actual_total_pages - total_pages_original_from_preview) > 1:
             logging.warning(f"Contagem de páginas divergente para {file_path}. Prévia: {total_pages_original_from_preview}, Atual: {actual_total_pages}")
             if not messagebox.askyesno("Aviso de Alteração", f"O número de páginas em {os.path.basename(file_path)} mudou de {total_pages_original_from_preview} para {actual_total_pages} desde a prévia.\nDeseja continuar a seleção com o intervalo '{page_range_str}' aplicado ao novo total de páginas?"):
                 return
             indices_to_select = parse_page_range(page_range_str, actual_total_pages)
             if not indices_to_select:
                 messagebox.showerror(texts["error"], "O intervalo fornecido não é mais válido com o novo total de páginas.")
                 return

        output_pdf = PdfWriter()
        progress_var.set(0)
        progress_bar["maximum"] = len(indices_to_select) 
        start_time = time.time()
        selected_page_count = 0

        for i, page_idx in enumerate(indices_to_select):
            if 0 <= page_idx < actual_total_pages:
                elapsed = time.time() - start_time
                eta = (elapsed / (i + 1)) * (len(indices_to_select) - (i + 1)) if i > 0 else 0
                status_label.config(text=texts["processing_page"].format(current=i + 1, total=len(indices_to_select), eta=eta))
                try:
                    output_pdf.add_page(input_pdf.pages[page_idx])
                    selected_page_count += 1
                except Exception as page_err: # Erro mais geral para renderização de página
                    logging.error(f"Erro ao adicionar página selecionada {page_idx+1} de {file_path}: {page_err}")
                
                progress_var.set(i + 1)
                animate_progress_bar()
                root.update_idletasks()
            else:
                logging.warning(f"Índice de página {page_idx+1} para seleção está fora do novo total de páginas ({actual_total_pages}). Pulando.")

        if selected_page_count == 0:
             messagebox.showwarning("Resultado Vazio", "Nenhuma página foi selecionada ou pôde ser adicionada. Nenhum arquivo foi gerado.")
             status_label.config(text="")
             progress_var.set(0)
             return

        output_filename = filedialog.asksaveasfilename(
            title="Salvar PDF com páginas selecionadas como...",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"{os.path.splitext(os.path.basename(file_path))[0]}_selecionado.pdf")

        if not output_filename:
            status_label.config(text="Operação cancelada.")
            progress_var.set(0)
            return

        with open(output_filename, "wb") as out_file:
            output_pdf.write(out_file)

        status_label.config(text="")
        logging.info(f"Páginas selecionadas com sucesso. Arquivo salvo como {output_filename}")
        messagebox.showinfo(texts["success"], f"Páginas selecionadas com sucesso!\nArquivo salvo em: {output_filename}")

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Arquivo não encontrado em select_pages_action: {file_path}")
        messagebox.showerror(texts["error"], f"Arquivo não encontrado: {os.path.basename(file_path)}")
    except Exception as e:
        status_label.config(text="Erro na operação.")
        logging.error(f"Erro em select_pages_action para {file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao selecionar páginas: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        if input_pdf_stream and not input_pdf_stream.closed:
            input_pdf_stream.close()

def add_pages_action(file_path_orig, file_path_second, insert_point_str, total_pages_orig_preview, total_pages_second_preview):
    """Acrescenta páginas de um PDF em outro em um ponto específico."""
    logging.info(f"Iniciando acréscimo. Original: {file_path_orig}, Adicional: {file_path_second}, Ponto: '{insert_point_str}'")
    try:
        insert_after_page_idx_user = int(insert_point_str) # Ponto de inserção baseado em 1, convertemos para 0-based
        if not (0 <= insert_after_page_idx_user <= total_pages_orig_preview): # 0 para início, total_pages para fim
            messagebox.showerror(texts["error"], texts["invalid_insert_point"].format(total=total_pages_orig_preview))
            return
    except ValueError:
        messagebox.showerror(texts["error"], "Ponto de inserção deve ser um número.")
        return

    pdf_orig_stream = None
    pdf_second_stream = None
    try:
        pdf_orig_stream = open(file_path_orig, "rb")
        pdf_second_stream = open(file_path_second, "rb")
        pdf_orig_reader = PdfReader(pdf_orig_stream)
        pdf_second_reader = PdfReader(pdf_second_stream)
        
        actual_total_orig = len(pdf_orig_reader.pages)
        actual_total_second = len(pdf_second_reader.pages)

        if actual_total_second == 0:
            messagebox.showinfo("PDF Adicional Vazio", "O PDF adicional está vazio. Nenhuma página será acrescentada.")
            return

        # Ajusta o ponto de inserção caso o PDF original tenha tido suas páginas alteradas
        if insert_after_page_idx_user > actual_total_orig:
            logging.warning(f"Ponto de inserção {insert_after_page_idx_user} excede novo total de páginas {actual_total_orig} do original. Inserindo no final.")
            insert_after_page_idx_user = actual_total_orig # Insere no final se o ponto original for maior

        output_pdf = PdfWriter()
        total_pages_to_process = actual_total_orig + actual_total_second
        progress_var.set(0)
        progress_bar["maximum"] = total_pages_to_process
        start_time = time.time()
        processed_pages_count = 0

        status_label.config(text="Preparando para acrescentar páginas...")
        root.update_idletasks()

        # Adiciona páginas do PDF original até o ponto de inserção
        for i in range(insert_after_page_idx_user):
            output_pdf.add_page(pdf_orig_reader.pages[i])
            processed_pages_count += 1
            if processed_pages_count % 20 == 0: progress_var.set(processed_pages_count); root.update_idletasks()

        # Adiciona páginas do segundo PDF
        num_added_from_second = 0
        for i in range(actual_total_second):
            output_pdf.add_page(pdf_second_reader.pages[i])
            processed_pages_count += 1
            num_added_from_second +=1
            elapsed = time.time() - start_time
            eta = (elapsed / processed_pages_count) * (total_pages_to_process - processed_pages_count) if processed_pages_count > 0 else 0
            status_label.config(text=texts["processing_page"].format(current=processed_pages_count, total=total_pages_to_process, eta=eta))
            progress_var.set(processed_pages_count)
            animate_progress_bar()
            root.update_idletasks()

        # Adiciona o restante das páginas do PDF original
        for i in range(insert_after_page_idx_user, actual_total_orig):
            output_pdf.add_page(pdf_orig_reader.pages[i])
            processed_pages_count += 1
            if processed_pages_count % 20 == 0: progress_var.set(processed_pages_count); root.update_idletasks()
        
        progress_var.set(total_pages_to_process)
        status_label.config(text="Salvando PDF final...")
        root.update_idletasks()

        output_filename = filedialog.asksaveasfilename(
            title="Salvar PDF com páginas acrescentadas como...",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"{os.path.splitext(os.path.basename(file_path_orig))[0]}_acrescentado.pdf")

        if not output_filename:
            status_label.config(text="Operação cancelada.")
            return

        with open(output_filename, "wb") as out_file:
            output_pdf.write(out_file)

        status_label.config(text="")
        logging.info(f"Páginas acrescentadas com sucesso. Arquivo salvo como {output_filename}")
        messagebox.showinfo(texts["success"], f"Páginas acrescentada(s) com sucesso!\nArquivo salvo em: {output_filename}")

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Um dos arquivos não foi encontrado em add_pages_action: {file_path_orig} ou {file_path_second}")
        messagebox.showerror(texts["error"], f"Um dos arquivos PDF não foi encontrado.")
    except Exception as e:
        status_label.config(text="Erro na operação.")
        logging.error(f"Erro em add_pages_action: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao acrescentar páginas: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        if pdf_orig_stream and not pdf_orig_stream.closed:
            pdf_orig_stream.close()
        if pdf_second_stream and not pdf_second_stream.closed:
            pdf_second_stream.close()

def replace_pages_action(file_path_orig, file_path_replace, page_range_str, total_pages_orig_preview, total_pages_replace_preview):
    """Substitui um intervalo de páginas de um PDF por páginas de outro PDF."""
    logging.info(f"Iniciando substituição. Original: {file_path_orig}, Substituto: {file_path_replace}, Range: '{page_range_str}'")

    indices_to_replace_orig = parse_page_range(page_range_str, total_pages_orig_preview)
    if not indices_to_replace_orig and page_range_str.strip():
        messagebox.showerror(texts["error"], texts["invalid_range"].format(total=total_pages_orig_preview) + "\nNenhuma página válida no intervalo para substituir.")
        return
    if not indices_to_replace_orig :
        messagebox.showinfo("Nenhuma Ação", "Nenhum intervalo de páginas válido fornecido para substituição.")
        return
    
    num_pages_to_replace_in_orig = len(indices_to_replace_orig)
    
    pdf_orig_stream = None
    pdf_replace_stream = None
    try:
        pdf_orig_stream = open(file_path_orig, "rb")
        pdf_replace_stream = open(file_path_replace, "rb")
        pdf_orig_reader = PdfReader(pdf_orig_stream)
        pdf_replace_reader = PdfReader(pdf_replace_stream)

        actual_total_orig = len(pdf_orig_reader.pages)
        actual_total_replace = len(pdf_replace_reader.pages)
        
        if actual_total_replace == 0:
            messagebox.showerror(texts["error"], "O PDF de substituição está vazio. Não é possível substituir páginas.")
            return

        # Verifica se as páginas do range original ainda existem
        valid_indices_in_actual_orig = [idx for idx in indices_to_replace_orig if 0 <= idx < actual_total_orig]
        if len(valid_indices_in_actual_orig) != num_pages_to_replace_in_orig:
            logging.warning(f"Algumas páginas do range '{page_range_str}' não existem mais no PDF original. Apenas páginas válidas serão consideradas.")
            if not valid_indices_in_actual_orig:
                messagebox.showerror(texts["error"], "Nenhuma das páginas do intervalo original existe mais no arquivo. Operação cancelada.")
                return
            num_pages_to_replace_in_orig = len(valid_indices_in_actual_orig) # Atualiza para o número real de páginas a substituir

        # Define quantas páginas do PDF substituto serão usadas
        num_replacement_pages_to_use = min(num_pages_to_replace_in_orig, actual_total_replace)
        if num_pages_to_replace_in_orig > actual_total_replace:
            # Aviso se o PDF substituto tem menos páginas que o intervalo a ser substituído
            if not messagebox.askyesno("Aviso de Substituição Parcial",
                                       f"O intervalo original tem {num_pages_to_replace_in_orig} página(s), "
                                       f"mas o PDF de substituição tem apenas {actual_total_replace} página(s).\n"
                                       f"Apenas as primeiras {actual_total_replace} página(s) do intervalo original "
                                       f"serão substituídas com as páginas do PDF de substituição. As restantes do intervalo original serão mantidas. Deseja continuar?"):
                return
        
        output_pdf = PdfWriter()
        progress_var.set(0)
        progress_bar["maximum"] = actual_total_orig # O progresso é baseado no total de páginas do original
        start_time = time.time()
        
        replace_pdf_page_counter = 0 # Contador para as páginas do PDF de substituição
        pages_actually_replaced_count = 0 # Contador de páginas realmente substituídas

        for i in range(actual_total_orig):
            elapsed = time.time() - start_time
            eta = (elapsed / (i + 1)) * (actual_total_orig - (i + 1)) if i > 0 else 0
            status_label.config(text=texts["processing_page"].format(current=i + 1, total=actual_total_orig, eta=eta))

            # Se a página atual é uma das que devem ser substituídas E ainda há páginas no PDF substituto
            if i in valid_indices_in_actual_orig and replace_pdf_page_counter < num_replacement_pages_to_use:
                output_pdf.add_page(pdf_replace_reader.pages[replace_pdf_page_counter])
                replace_pdf_page_counter += 1
                pages_actually_replaced_count +=1
            else:
                # Caso contrário, mantém a página original
                output_pdf.add_page(pdf_orig_reader.pages[i])
            
            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()

        output_filename = filedialog.asksaveasfilename(
            title="Salvar PDF com páginas substituídas como...",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"{os.path.splitext(os.path.basename(file_path_orig))[0]}_substituido.pdf")

        if not output_filename:
            status_label.config(text="Operação cancelada.")
            return

        with open(output_filename, "wb") as out_file:
            output_pdf.write(out_file)

        status_label.config(text="")
        logging.info(f"{pages_actually_replaced_count} página(s) substituída(s) com sucesso. Arquivo salvo como {output_filename}")
        messagebox.showinfo(texts["success"], f"{pages_actually_replaced_count} página(s) substituída(s) com sucesso!\nArquivo salvo em: {output_filename}")

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Um dos arquivos não foi encontrado em replace_pages_action: {file_path_orig} ou {file_path_replace}")
        messagebox.showerror(texts["error"], f"Um dos arquivos PDF não foi encontrado.")
    except Exception as e:
        status_label.config(text="Erro na operação.")
        logging.error(f"Erro em replace_pages_action: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao substituir páginas: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        if pdf_orig_stream and not pdf_orig_stream.closed:
            pdf_orig_stream.close()
        if pdf_replace_stream and not pdf_replace_stream.closed:
            pdf_replace_stream.close()

# --- Funções de Split, Merge (Mantêm fluxo no options_frame ou direto) ---
def on_split_button():
    """Prepara a interface no options_frame para as opções de divisão de PDF."""
    file_path = select_file(title="Selecione o PDF para Separar")
    if not file_path: return

    try:
        pdf_reader_stream = open(file_path, "rb") # Abre o stream para leitura
        pdf_reader = PdfReader(pdf_reader_stream)
        total_pages = len(pdf_reader.pages)
        pdf_reader_stream.close() # Fecha o stream imediatamente após obter a contagem de páginas
    except Exception as e:
        logging.error(f"Erro ao ler PDF para split: {file_path} - {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao abrir ou ler PDF: {os.path.basename(file_path)}\n({e})")
        return
    
    if total_pages < 2 : # PDFs com 0 ou 1 página não podem ser divididos
        messagebox.showwarning("Divisão Inválida", "O PDF precisa ter pelo menos 2 páginas para ser dividido.")
        return

    clear_options_frame()
    description = get_description_for_function("split_pdf") 

    # Fundo claro, texto escuro
    ttkb.Label(options_frame, text=texts["custom_parts_prompt"], font=("Georgia", 12, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(10,5)) 
    # Fundo claro, texto marrom escuro
    ttkb.Label(options_frame, text=description, background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, wraplength=380, justify="center", font=("Verdana", 9)).pack(pady=(0,15)) 

    button_width = 22
    # Estilo dos botões: texto Noite Estrelada, fundo Luz Pura da Lua, borda Terracota
    btn_half = ttkb.Button(options_frame, text=texts["split_half"], command=lambda: split_pdf(file_path, 2, total_pages), width=button_width, style='GuaraOutlineCategory.TButton')
    btn_half.pack(pady=4, fill="x", padx=20)
    create_tooltip(btn_half, texts["tooltip_split"])

    btn_third = ttkb.Button(options_frame, text=texts["split_third"], command=lambda: split_pdf(file_path, 3, total_pages), width=button_width, style='GuaraOutlineCategory.TButton')
    btn_third.pack(pady=4, fill="x", padx=20)
    create_tooltip(btn_third, texts["tooltip_split"])

    btn_quarter = ttkb.Button(options_frame, text=texts["split_quarter"], command=lambda: split_pdf(file_path, 4, total_pages), width=button_width, style='GuaraOutlineCategory.TButton')
    btn_quarter.pack(pady=4, fill="x", padx=20)
    create_tooltip(btn_quarter, texts["tooltip_split"])

    btn_custom = ttkb.Button(options_frame, text=texts["split_custom"], command=lambda: show_custom_split(file_path, total_pages), width=button_width, style='GuaraOutlineCategory.TButton')
    btn_custom.pack(pady=(10,4), fill="x", padx=20)
    create_tooltip(btn_custom, texts["tooltip_split"])

def split_pdf(file_path, num_parts, total_pages_from_caller):
    """Divide o PDF no número especificado de partes."""
    logging.info(f"Iniciando divisão de {file_path} em {num_parts} partes.")
    
    input_pdf_stream = None # Inicializa stream como None para garantir fechamento
    try:
        input_pdf_stream = open(file_path, "rb")
        input_pdf = PdfReader(input_pdf_stream)
        actual_total_pages = len(input_pdf.pages)

        if actual_total_pages < num_parts: # Verifica se o número de partes é viável
             messagebox.showerror(texts["error"], f"Não é possível dividir {actual_total_pages} páginas em {num_parts} partes.")
             return

        # Aviso se a contagem de páginas mudou desde a seleção
        if actual_total_pages != total_pages_from_caller:
            logging.warning(f"Contagem de páginas divergente em split_pdf para {file_path}. Chamador: {total_pages_from_caller}, Atual: {actual_total_pages}")
            if not messagebox.askyesno("Aviso de Alteração", f"O PDF {os.path.basename(file_path)} tem {actual_total_pages} páginas (diferente da contagem inicial de {total_pages_from_caller}). Deseja continuar a divisão?"):
                return
        
        output_base_path = os.path.dirname(file_path)
        original_filename_no_ext = os.path.splitext(os.path.basename(file_path))[0]
        
        pages_per_part = math.ceil(actual_total_pages / num_parts) # Garante que todas as páginas sejam incluídas

        progress_var.set(0)
        progress_bar["maximum"] = num_parts # Progresso por parte criada
        start_time = time.time()
        current_page_start_index = 0
        parts_created_count = 0

        for i in range(num_parts):
            if current_page_start_index >= actual_total_pages: break # Sai se não há mais páginas

            elapsed = time.time() - start_time
            eta = (elapsed / (i + 1)) * (num_parts - (i + 1)) if i > 0 else 0
            status_label.config(text=texts["processing_page"].format(current=i + 1, total=num_parts, eta=eta))
            
            output_pdf_part = PdfWriter()
            current_page_end_index = min(current_page_start_index + pages_per_part, actual_total_pages)
            
            for page_num in range(current_page_start_index, current_page_end_index):
                output_pdf_part.add_page(input_pdf.pages[page_num])
            
            if not output_pdf_part.pages: # Se por algum motivo a parte ficou vazia, pula
                logging.warning(f"Parte {i+1} para {file_path} resultou vazia. Pulando.")
                continue

            part_filename = os.path.join(output_base_path, f"{original_filename_no_ext}_parte_{i+1}.pdf")
            with open(part_filename, "wb") as out_file:
                output_pdf_part.write(out_file)
            parts_created_count +=1
            
            current_page_start_index = current_page_end_index # Atualiza o índice para a próxima parte
            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()

        status_label.config(text="")
        logging.info(f"PDF {file_path} dividido com sucesso em {parts_created_count} partes.")
        messagebox.showinfo(texts["success"], f"PDF dividido em {parts_created_count} partes!\nSalvo em: {output_base_path}")

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Arquivo não encontrado em split_pdf: {file_path}")
        messagebox.showerror(texts["error"], f"Arquivo não encontrado: {os.path.basename(file_path)}")
    except Exception as e:
        status_label.config(text="Erro na divisão.")
        logging.error(f"Erro em split_pdf para {file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao dividir PDF: {str(e)}\nConsulte o log.")
    finally:
        if input_pdf_stream and not input_pdf_stream.closed: # Garante que o stream seja fechado
            input_pdf_stream.close()
        progress_var.set(0)
        clear_options_frame() # Limpa as opções após a conclusão
        show_welcome_panel() # Volta ao painel de boas-vindas

def show_custom_split(file_path, total_pages):
    """Mostra o input para divisão customizada no options_frame."""
    clear_options_frame()
    description = get_description_for_function("split_pdf") 

    # Fundo claro, texto escuro
    ttkb.Label(options_frame, text=texts["custom_parts_prompt"], font=("Georgia", 12, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(10,5)) 
    # Fundo claro, texto marrom escuro
    ttkb.Label(options_frame, text=f"O PDF selecionado tem {total_pages} páginas.", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, font=("Verdana", 9)).pack(pady=(0,10)) 
    
    # Entry com fundo claro e texto escuro
    entry_parts = ttkb.Entry(options_frame, width=10, font=("Verdana", 10), fieldbackground=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY, bordercolor=COLOR_TERRACOTTA) 
    entry_parts.pack(pady=5)
    entry_parts.focus_set()

    def _confirm_custom_split():
        parts_str = entry_parts.get()
        if validate_parts(parts_str, total_pages):
            split_pdf(file_path, int(parts_str), total_pages) # Passa total_pages para revalidação
        else:
            messagebox.showerror(texts["error"], texts["invalid_parts"] + f"\n(Deve ser entre 2 e {total_pages})")
            entry_parts.focus_set()
            entry_parts.select_range(0, 'end') # Seleciona o texto para facilitar a correção

    # Botão sólido terracota, texto noite estrelada
    confirm_button = ttkb.Button(options_frame, text="Confirmar Divisão Customizada", command=_confirm_custom_split, width=30, background=COLOR_TERRACOTTA, foreground=COLOR_NIGHT_SKY, font=("Verdana", 10, "bold")) 
    confirm_button.pack(pady=10)
    entry_parts.bind("<Return>", lambda e: _confirm_custom_split())

def on_merge_button():
    """Mescla múltiplos arquivos PDF selecionados pelo usuário."""
    clear_options_frame() # Limpa o painel de opções antes de abrir a seleção
    file_paths = select_files(title="Selecione os PDFs para Mesclar (ordem de seleção importa)")
    
    if not file_paths or len(file_paths) < 2:
        messagebox.showwarning("Seleção Insuficiente", "Por favor, selecione pelo menos dois arquivos PDF para mesclar.")
        show_welcome_panel() # Volta ao painel de boas-vindas
        return

    output_filename = filedialog.asksaveasfilename(
        title="Salvar PDF Mesclado como...",
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        initialfile="PDF_Mesclado.pdf")
    
    if not output_filename:
        status_label.config(text="Operação de mesclagem cancelada.")
        return

    input_streams = [] # Lista para manter referências a streams abertos
    try:
        merged_pdf_writer = PdfWriter()
        estimated_total_pages = 0
        
        # Primeira passada para estimar o total de páginas (para a barra de progresso)
        for fp_scan in file_paths:
            try:
                temp_stream = open(fp_scan, "rb")
                reader = PdfReader(temp_stream)
                estimated_total_pages += len(reader.pages)
                temp_stream.close() # Fecha o stream de leitura temporária
            except Exception as e_scan:
                logging.warning(f"Não foi possível ler {fp_scan} para contagem de páginas na mesclagem: {e_scan}")
        
        progress_var.set(0)
        progress_bar["maximum"] = estimated_total_pages if estimated_total_pages > 0 else 1 # Evita divisão por zero
        start_time = time.time()
        current_page_processed_count = 0

        for i, file_path_to_merge in enumerate(file_paths):
            status_label.config(text=f"Processando arquivo {i+1}/{len(file_paths)}: {os.path.basename(file_path_to_merge)}...")
            root.update_idletasks()
            try:
                current_stream = open(file_path_to_merge, "rb")
                input_streams.append(current_stream) # Adiciona à lista para fechar depois
                input_pdf = PdfReader(current_stream)
                for page_num_in_file in range(len(input_pdf.pages)):
                    merged_pdf_writer.add_page(input_pdf.pages[page_num_in_file])
                    current_page_processed_count +=1
                    if estimated_total_pages > 0 : # Evita divisão por zero se o total for 0
                        elapsed = time.time() - start_time
                        eta = (elapsed / current_page_processed_count) * (estimated_total_pages - current_page_processed_count) if current_page_processed_count > 0 else 0
                        if current_page_processed_count % 5 == 0 or current_page_processed_count == estimated_total_pages : # Atualiza a cada 5 páginas ou no final
                            status_label.config(text=texts["processing_page"].format(current=current_page_processed_count, total=estimated_total_pages, eta=eta))
                    progress_var.set(current_page_processed_count)
                    if current_page_processed_count % 10 == 0: # Anima a cada 10 páginas
                        animate_progress_bar()
                        root.update_idletasks()

            except Exception as e_merge_file:
                logging.error(f"Erro ao mesclar o arquivo {file_path_to_merge}: {e_merge_file}")
                messagebox.showwarning("Erro de Arquivo", f"Não foi possível processar o arquivo:\n{os.path.basename(file_path_to_merge)}\nSerá pulado.")
                continue # Continua para o próximo arquivo mesmo se um falhar

        if not merged_pdf_writer.pages: # Verifica se alguma página foi realmente adicionada
             messagebox.showerror(texts["error"], "Nenhuma página foi adicionada ao PDF mesclado. Verifique os arquivos de entrada.")
             status_label.config(text="")
             progress_var.set(0)
             return

        with open(output_filename, "wb") as out_file:
            merged_pdf_writer.write(out_file)
        
        status_label.config(text="")
        logging.info(f"PDFs mesclados com sucesso em {output_filename}")
        messagebox.showinfo(texts["success"], f"PDFs mesclados com sucesso!\nSalvo em: {output_filename}")

    except Exception as e:
        status_label.config(text="Erro na mesclagem.")
        logging.error(f"Erro geral em on_merge_button: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao mesclar PDFs: {str(e)}\nConsulte o log.")
    finally:
        for stream in input_streams: # Garante que todos os streams abertos sejam fechados
            if not stream.closed:
                stream.close()
        progress_var.set(0)
        show_welcome_panel() # Volta ao painel de boas-vindas

# --- Funções de Conversão e Otimização ---
def choose_conversion(): 
    """Prepara a interface para a escolha de conversão PDF <-> JPG."""
    clear_options_frame()
    description = get_description_for_function("pdf_jpg")

    # Fundo claro, texto escuro
    ttkb.Label(options_frame, text=texts["choose_conversion"], font=("Georgia", 12, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(10,5)) 
    # Fundo claro, texto marrom escuro
    ttkb.Label(options_frame, text=description, background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, wraplength=380, justify="center", font=("Verdana", 9)).pack(pady=(0,15)) 
    
    button_width = 25
    # Estilo dos botões: texto Noite Estrelada, fundo Luz Pura da Lua, borda colorida
    pdf_to_jpg_btn = ttkb.Button(options_frame, text=texts["pdf_to_jpg"], 
                                 command=lambda: pdf_to_jpg(select_file(title="Selecione o PDF para converter para JPG")), 
                                 width=button_width, style='GuaraOutlineCategory.TButton')
    pdf_to_jpg_btn.pack(pady=4, fill="x", padx=20)
    create_tooltip(pdf_to_jpg_btn, texts["tooltip_pdf_jpg"])

    jpg_to_pdf_btn = ttkb.Button(options_frame, text=texts["jpg_to_pdf"], command=jpg_to_pdf, 
                                 width=button_width, style='GuaraOutlineCategory.TButton')
    jpg_to_pdf_btn.pack(pady=4, fill="x", padx=20)
    create_tooltip(jpg_to_pdf_btn, texts["tooltip_pdf_jpg"])

def pdf_to_jpg(pdf_file_path):
    """Converte um PDF em imagens JPG (uma imagem por página)."""
    if not pdf_file_path: return
    output_folder = filedialog.askdirectory(title="Selecione a pasta para salvar os JPGs")
    if not output_folder: return

    try:
        pdf_document = fitz.open(pdf_file_path)
        total_pages = len(pdf_document)
        if total_pages == 0:
            messagebox.showinfo("PDF Vazio", "O PDF selecionado não contém páginas.")
            pdf_document.close()
            return

        progress_var.set(0)
        progress_bar["maximum"] = total_pages
        start_time = time.time()
        for page_number in range(total_pages):
            elapsed = time.time() - start_time
            eta = (elapsed / (page_number + 1)) * (total_pages - (page_number + 1)) if page_number > 0 else 0
            status_label.config(text=texts["processing_page"].format(current=page_number + 1, total=total_pages, eta=eta))
            
            page = pdf_document.load_page(page_number)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # Renderiza com alta resolução (2x)
            output_filename = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(pdf_file_path))[0]}_pagina_{page_number + 1}.jpg")
            pix.save(output_filename) # Salva como JPG
            
            progress_var.set(page_number + 1)
            animate_progress_bar()
            root.update_idletasks()
            
        pdf_document.close()
        status_label.config(text="")
        logging.info(f"PDF {pdf_file_path} convertido para JPGs em {output_folder}")
        messagebox.showinfo(texts["success"], f"PDF convertido para JPG!\nImagens salvas em: {output_folder}")
    except Exception as e:
        status_label.config(text="Erro na conversão.")
        logging.error(f"Erro em pdf_to_jpg para {pdf_file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro na conversão para JPG: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        clear_options_frame()
        show_welcome_panel()

def jpg_to_pdf():
    """Converte múltiplos arquivos JPG/JPEG em um único PDF."""
    image_files = filedialog.askopenfilenames(
        title="Selecione os arquivos JPG/JPEG (ordem de seleção será a ordem no PDF)",
        filetypes=[("Image Files", "*.jpg *.jpeg")])
    if not image_files: return

    pdf_path = filedialog.asksaveasfilename(
        title="Salvar PDF como...",
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        initialfile="imagens_convertidas.pdf")
    if not pdf_path: return

    output_pdf = fitz.open() # Cria um novo documento PDF
    try:
        total_files = len(image_files)
        progress_var.set(0)
        progress_bar["maximum"] = total_files
        start_time = time.time()
        images_converted_count = 0

        for i, img_path in enumerate(image_files):
            elapsed = time.time() - start_time
            eta = (elapsed / (i + 1)) * (total_files - (i + 1)) if i > 0 else 0
            status_label.config(text=f"Processando imagem {i+1}/{total_files}: {os.path.basename(img_path)}... (ETA: {eta:.1f}s)")
            
            try:
                img_doc = fitz.open(img_path) # Abre a imagem como um "documento" fitz
                pdf_bytes = img_doc.convert_to_pdf() # Converte a imagem para bytes de PDF
                img_doc.close() # Fecha o documento da imagem
                
                img_pdf_page = fitz.open("pdf", pdf_bytes) # Abre esses bytes como um PDF temporário
                output_pdf.insert_pdf(img_pdf_page) # Insere no PDF de saída
                img_pdf_page.close() # Fecha o PDF temporário
                images_converted_count +=1
            except Exception as img_e:
                logging.error(f"Erro ao processar imagem {img_path} para PDF: {img_e}")
                messagebox.showwarning("Erro Imagem", f"Não foi possível converter {os.path.basename(img_path)}. Será pulada.")
                continue # Continua para a próxima imagem

            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()

        if images_converted_count > 0:
            output_pdf.save(pdf_path, garbage=3, deflate=True) # Salva o PDF otimizado
            status_label.config(text="")
            logging.info(f"{images_converted_count} imagem(ns) convertida(s) para PDF: {pdf_path}")
            messagebox.showinfo(texts["success"], f"{images_converted_count} imagem(ns) convertida(s) para PDF!\nSalvo em: {pdf_path}")
        else:
            status_label.config(text="Nenhuma imagem convertida.")
            messagebox.showwarning("Sem Conversão", "Nenhuma imagem pôde ser convertida para PDF.")
            
    except Exception as e:
        status_label.config(text="Erro na conversão.")
        logging.error(f"Erro em jpg_to_pdf: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro na conversão para PDF: {str(e)}\nConsulte o log.")
    finally:
        if output_pdf and not output_pdf.is_closed: # Garante que o documento PDF de saída seja fechado
            output_pdf.close()
        progress_var.set(0)
        clear_options_frame()
        show_welcome_panel()

def on_convert_button(): 
    """Prepara a interface para a escolha de conversão de PDF para Word ou Excel."""
    clear_options_frame()
    description = get_description_for_function("convert_pdf") 
    
    # Fundo claro, texto escuro
    ttkb.Label(options_frame, text=texts["choose_output"], font=("Georgia", 12, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(10,5)) 
    # Fundo claro, texto marrom escuro
    ttkb.Label(options_frame, text=description, background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, wraplength=380, justify="center", font=("Verdana", 9)).pack(pady=(0,15)) 

    button_width = 25
    # Estilo dos botões: texto Noite Estrelada, fundo Luz Pura da Lua, borda colorida
    word_btn = ttkb.Button(options_frame, text=texts["word"], 
                            command=lambda: save_text_as_word(select_file(title="Selecione o PDF para converter para Word")), 
                            width=button_width, style='GuaraOutlineCategory.TButton')
    word_btn.pack(pady=4, fill="x", padx=20)
    create_tooltip(word_btn, texts["tooltip_convert"])

    excel_btn = ttkb.Button(options_frame, text=texts["excel"], 
                             command=lambda: save_text_as_excel(select_file(title="Selecione o PDF para converter para Excel")), 
                             width=button_width, style='GuaraOutlineCategory.TButton')
    excel_btn.pack(pady=4, fill="x", padx=20)
    create_tooltip(excel_btn, texts["tooltip_convert"])

def save_text_as_word(pdf_file_path):
    """Extrai texto de um PDF e salva como documento Word (.docx)."""
    if not pdf_file_path: return
    output_file = filedialog.asksaveasfilename(
        title="Salvar como Word (.docx)",
        defaultextension=".docx", 
        filetypes=[("Word Document", "*.docx")],
        initialfile=f"{os.path.splitext(os.path.basename(pdf_file_path))[0]}.docx")
    if not output_file: return

    try:
        pdf_document = fitz.open(pdf_file_path)
        doc = Document() # Cria um novo documento Word
        status_label.config(text="Extraindo texto para Word...")
        progress_var.set(0)
        progress_bar["maximum"] = len(pdf_document)
        
        full_text = ""
        for i, page in enumerate(pdf_document):
            status_label.config(text=f"Processando página {i+1}/{len(pdf_document)} para Word...")
            full_text += page.get_text("text") + "\n" # Extrai texto e adiciona quebra de linha
            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()
            
        doc.add_paragraph(full_text) # Adiciona todo o texto extraído como um parágrafo
        pdf_document.close()
        doc.save(output_file)
        status_label.config(text="")
        logging.info(f"PDF {pdf_file_path} convertido para Word: {output_file}")
        messagebox.showinfo(texts["success"], f"Convertido para Word!\nSalvo em: {output_file}")
    except Exception as e:
        status_label.config(text="Erro na conversão Word.")
        logging.error(f"Erro em save_text_as_word para {pdf_file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao converter para Word: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        clear_options_frame()
        show_welcome_panel()

def save_text_as_excel(pdf_file_path):
    """Extrai texto e tabelas de um PDF e salva como documento Excel (.xlsx)."""
    if not pdf_file_path: return
    output_file = filedialog.asksaveasfilename(
        title="Salvar como Excel (.xlsx)",
        defaultextension=".xlsx", 
        filetypes=[("Excel Workbook", "*.xlsx")],
        initialfile=f"{os.path.splitext(os.path.basename(pdf_file_path))[0]}.xlsx")
    if not output_file: return

    try:
        pdf_document = fitz.open(pdf_file_path)
        status_label.config(text="Extraindo tabelas/texto para Excel...")
        progress_var.set(0)
        progress_bar["maximum"] = len(pdf_document)
        
        all_page_data_dfs = [] # Lista para armazenar DataFrames de cada página
        for i, page in enumerate(pdf_document):
            status_label.config(text=f"Analisando página {i+1}/{len(pdf_document)} para Excel...")
            
            tables = page.find_tables()
            if tables.tables: # Se tabelas forem detectadas
                for table_obj in tables:
                    df = table_obj.to_pandas() # Converte a tabela para DataFrame
                    if not df.empty:
                        all_page_data_dfs.append(df)
            else: # Se nenhuma tabela for detectada, tenta extrair blocos de texto como linhas
                blocks = page.get_text("blocks") # Extrai blocos de texto
                page_data_as_rows = []
                for b in blocks:
                    lines_in_block = b[4].strip().split('\n') # Aposição [4] contém o texto do bloco
                    for line_text in lines_in_block:
                        if line_text.strip(): # Adiciona apenas linhas não vazias
                            page_data_as_rows.append([line_text.strip()]) # Cada linha como uma nova linha no DataFrame
                if page_data_as_rows:
                    df = pd.DataFrame(page_data_as_rows)
                    all_page_data_dfs.append(df)

            progress_var.set(i + 1)
            animate_progress_bar()
            root.update_idletasks()
            
        pdf_document.close()

        if not all_page_data_dfs:
            messagebox.showwarning("Sem Conteúdo", "Nenhuma tabela estruturada ou bloco de texto significativo pôde ser extraído para Excel.")
            status_label.config(text="")
            return

        with pd.ExcelWriter(output_file) as writer:
            sheet_count = 0
            for df_idx, df_item in enumerate(all_page_data_dfs):
                sheet_name = f'Dados_Extraidos_{sheet_count+1}'
                actual_sheet_name = sheet_name
                k=1
                while actual_sheet_name in writer.sheets: # Garante nome de aba único
                    actual_sheet_name = f"{sheet_name}_{k}"
                    k+=1
                df_item.to_excel(writer, sheet_name=actual_sheet_name, index=False, header=False) # Sem índice e sem cabeçalho padrão
                if (df_idx + 1) < len(all_page_data_dfs) : # Incrementa apenas se houver mais dataframes
                     sheet_count +=1
        
        status_label.config(text="")
        logging.info(f"PDF {pdf_file_path} convertido para Excel: {output_file}")
        messagebox.showinfo(texts["success"], f"Conteúdo extraído para Excel!\nSalvo em: {output_file}\n(Nota: A extração pode variar conforme a estrutura do PDF)")
    except Exception as e:
        status_label.config(text="Erro na conversão Excel.")
        logging.error(f"Erro em save_text_as_excel para {pdf_file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao converter para Excel: {str(e)}\nConsulte o log.")
    finally:
        progress_var.set(0)
        clear_options_frame()
        show_welcome_panel()

def anonymize_pdf():
    """Remove metadados de um PDF para conformidade com a LGPD."""
    file_path = select_file(title="Selecione o PDF para Anonimizar Metadados")
    if not file_path: return

    output_filename = filedialog.asksaveasfilename(
        title="Salvar PDF Anonimizado como...",
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        initialfile=f"{os.path.splitext(os.path.basename(file_path))[0]}_anonimizado.pdf")
    if not output_filename: return

    try:
        pdf_doc = fitz.open(file_path)
        current_metadata = pdf_doc.metadata
        
        # Cria um novo dicionário de metadados com todos os valores como None
        new_metadata = { k: None for k in current_metadata.keys() if current_metadata[k] is not None} # remove apenas os que não são None
        logging.info(f"Anonimizando metadata: {new_metadata.keys()}")
        if new_metadata: # Apenas tenta setar se houver algo para remover
             pdf_doc.set_metadata(new_metadata)
        else:
             logging.info("Nenhum metadado encontrado para remover.")

        # Salva o PDF com otimização, garbage collection e limpeza (para remover objetos órfãos)
        pdf_doc.save(output_filename, garbage=4, deflate=True, clean=True) 
        pdf_doc.close()
        
        logging.info(f"Metadados do PDF {file_path} anonimizados e salvos em {output_filename}")
        messagebox.showinfo(texts["success"], f"Metadados do PDF removidos/limpos.\nSalvo em: {output_filename}")
    except Exception as e:
        logging.error(f"Erro em anonymize_pdf para {file_path}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao anonimizar PDF: {str(e)}\nConsulte o log.")
    finally:
        clear_options_frame() # Limpa as opções após a conclusão
        show_welcome_panel()

def compact_pdf(): 
    """Prepara a interface para a escolha do nível de compactação de PDF."""
    file_path = select_file(title="Selecione o PDF para Compactar")
    if not file_path: return

    clear_options_frame()
    description = get_description_for_function("compress_pdf")

    # Fundo claro, texto escuro
    ttkb.Label(options_frame, text=texts["compress_level"], font=("Georgia", 12, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=(10,5)) 
    # Fundo claro, texto marrom escuro
    ttkb.Label(options_frame, text=description, background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, wraplength=380, justify="center", font=("Verdana", 9)).pack(pady=(0,15)) 

    button_width = 25
    # Estilo dos botões: texto Noite Estrelada, fundo Luz Pura da Lua, borda colorida
    btn_light = ttkb.Button(options_frame, text=texts["light"],
                             command=lambda: compact_pdf_with_level(file_path, "leve"),
                             width=button_width, style='GuaraOutlineSuccessCategory.TButton')
    btn_light.pack(pady=4, fill="x", padx=20)
    
    btn_mod = ttkb.Button(options_frame, text=texts["moderate"],
                            command=lambda: compact_pdf_with_level(file_path, "moderada"),
                            width=button_width, style='GuaraOutlineWarningCategory.TButton')
    btn_mod.pack(pady=4, fill="x", padx=20)
    
    btn_agg = ttkb.Button(options_frame, text=texts["aggressive"],
                            command=lambda: compact_pdf_with_level(file_path, "agressiva"),
                            width=button_width, style='GuaraOutlineDangerCategory.TButton')
    btn_agg.pack(pady=4, fill="x", padx=20)

def process_page(page_num, pdf_path_for_worker, temp_dir_for_worker, dpi_for_worker, image_quality_for_worker):
    """Processa uma única página para compactação (executado em thread)."""
    img_path_result = None
    original_width, original_height = None, None
    try:
        pdf_doc_worker = fitz.open(pdf_path_for_worker) # Abre o documento localmente para o thread
        if 0 <= page_num < len(pdf_doc_worker):
            page = pdf_doc_worker.load_page(page_num)
            original_width, original_height = page.rect.width, page.rect.height # Dimensões originais para criar nova página
            
            pix = page.get_pixmap(dpi=dpi_for_worker) # Renderiza a página com a DPI especificada
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            img_path_result = os.path.join(temp_dir_for_worker, f"page_{page_num}.jpg")
            img.save(img_path_result, format="JPEG", quality=image_quality_for_worker, optimize=True)
            img.close()
        else:
             logging.warning(f"Índice de página inválido {page_num} para {pdf_path_for_worker} em process_page.")
        pdf_doc_worker.close() # Fecha o documento após o uso pelo thread
    except Exception as e:
        logging.error(f"Erro ao processar página {page_num} de {pdf_path_for_worker} para compactação: {str(e)}")
        img_path_result = None # Sinaliza que houve erro
    return page_num, img_path_result, original_width, original_height

def compact_pdf_with_level(file_path, level):
    """Compacta o PDF convertendo páginas para imagens JPEG com o nível de qualidade especificado."""
    logging.info(f"Iniciando compactação de {file_path} com nível {level}.")
    
    output_filename = filedialog.asksaveasfilename(
        title=f"Salvar PDF Compactado ({level}) como...",
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        initialfile=f"{os.path.splitext(os.path.basename(file_path))[0]}_compactado_{level}.pdf")
    if not output_filename: return

    try:
        pdf_main_doc_check = fitz.open(file_path)
        total_pages = len(pdf_main_doc_check)
        pdf_main_doc_check.close() # Fecha para liberar o arquivo para os threads

        if total_pages == 0:
            messagebox.showinfo("PDF Vazio", "O PDF selecionado não contém páginas para compactar.")
            return

        new_pdf_output = fitz.open() # Novo documento PDF para as páginas compactadas
        quality_map = {
            "leve": {"dpi": 150, "image_quality": 85},    # DPI mais alta, qualidade de imagem boa
            "moderada": {"dpi": 100, "image_quality": 65},   # DPI média, qualidade média
            "agressiva": {"dpi": 72, "image_quality": 40}   # DPI de tela, qualidade de imagem mais baixa
        }
        settings = quality_map[level]
        dpi = settings["dpi"]
        image_quality = settings["image_quality"] 
        
        progress_var.set(0)
        progress_bar["maximum"] = total_pages
        start_time = time.time()
        
        page_processing_results = [None] * total_pages # Lista para armazenar resultados na ordem correta

        with tempfile.TemporaryDirectory() as temp_dir: # Cria um diretório temporário para as imagens
            with ThreadPoolExecutor(max_workers=os.cpu_count() or 1) as executor: # Usa todos os núcleos da CPU
                future_to_page = {
                    executor.submit(process_page, i, file_path, temp_dir, dpi, image_quality): i
                    for i in range(total_pages)
                }
                
                processed_count = 0
                for future in as_completed(future_to_page.keys()): # Itera conforme as tarefas são concluídas
                    page_num_completed = future_to_page[future]
                    try:
                        _, img_path, width, height = future.result()
                        if img_path and width is not None and height is not None: # Verifica se o processamento foi bem-sucedido
                            page_processing_results[page_num_completed] = (img_path, width, height)
                        else: # Loga se a imagem temporária não foi gerada
                            logging.warning(f"Falha ao processar página {page_num_completed+1} de {file_path}. Será pulada.")
                    except Exception as exc:
                        logging.error(f"Página {page_num_completed+1} gerou exceção durante compactação: {exc}")
                    
                    processed_count +=1
                    elapsed = time.time() - start_time
                    eta = (elapsed / processed_count) * (total_pages - processed_count) if processed_count > 0 else 0
                    status_label.config(text=texts["processing_page"].format(current=processed_count, total=total_pages, eta=eta))
                    progress_var.set(processed_count)
                    if processed_count % (total_pages // 10 or 1) == 0 : # Anima a barra de progresso a cada 10%
                         animate_progress_bar()
                         root.update_idletasks()
            
            status_label.config(text="Montando PDF compactado...")
            root.update_idletasks()
            actual_pages_added = 0
            for i in range(total_pages):
                result = page_processing_results[i]
                if result:
                    img_path, width, height = result
                    if os.path.exists(img_path): # Confirma que a imagem existe no disco
                        page_output = new_pdf_output.new_page(width=width, height=height) # Cria uma nova página com as dimensões originais
                        page_output.insert_image(page_output.rect, filename=img_path) # Insere a imagem compactada
                        actual_pages_added += 1
                    else:
                        logging.warning(f"Imagem temporária não encontrada para página {i+1}: {img_path}")
                else: # Se o resultado para a página é None, significa que houve um erro no processamento do thread
                     logging.info(f"Página {i+1} foi pulada na montagem final devido a erro no processamento.")

        if actual_pages_added == 0:
             messagebox.showerror(texts["error"],"Nenhuma página pôde ser processada com sucesso para compactação.")
             status_label.config(text="")
             progress_var.set(0)
             if new_pdf_output and not new_pdf_output.is_closed: new_pdf_output.close()
             return

        new_pdf_output.save(output_filename, garbage=4, deflate=True) # Salva o PDF otimizado
        old_size_mb = get_file_size_mb(file_path)
        new_size_mb = get_file_size_mb(output_filename)
        
        status_label.config(text="")
        logging.info(f"PDF {file_path} compactado com nível {level} e salvo como {output_filename}. Tamanho: {old_size_mb:.2f}MB -> {new_size_mb:.2f}MB")
        messagebox.showinfo(texts["success"], texts["compress_success"].format(old_size=old_size_mb, new_size=new_size_mb))

    except FileNotFoundError:
        status_label.config(text="Erro: Arquivo não encontrado.")
        logging.error(f"Arquivo não encontrado para compactação: {file_path}")
        messagebox.showerror(texts["error"], f"Arquivo não encontrado: {os.path.basename(file_path)}")
    except Exception as e:
        status_label.config(text="Erro na compactação.")
        logging.error(f"Erro em compact_pdf_with_level para {file_path}, nível {level}: {str(e)}")
        messagebox.showerror(texts["error"], f"Erro ao compactar PDF: {str(e)}\nConsulte o log.")
    finally:
        # Garante que o documento PDF de saída seja fechado
        if 'new_pdf_output' in locals() and hasattr(new_pdf_output, 'is_closed') and not new_pdf_output.is_closed:
            new_pdf_output.close()
        progress_var.set(0)
        clear_options_frame()
        show_welcome_panel()

# --- Funções de Ajuda e Animação ---
def show_help():
    """Exibe o painel de ajuda com informações sobre o software."""
    clear_options_frame() 
    # Título da ajuda com fundo Luz Pura da Lua e texto Noite Estrelada
    ttkb.Label(options_frame, text=texts["help"], font=("Georgia", 14, "bold"), background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY).pack(pady=10, fill="x")
    
    help_text_area = tk.Text(options_frame, wrap="word", font=("Verdana", 10), 
                             relief="flat", borderwidth=0, padx=10, pady=10,
                             bg=COLOR_MOON_LIGHT, # Fundo do texto da ajuda em Luz Pura da Lua
                             fg=COLOR_NIGHT_SKY) # Texto da ajuda em Noite Estrelada
    
    scrollbar_help = ttkb.Scrollbar(options_frame, orient="vertical", command=help_text_area.yview, bootstyle="round-info")
    # Customizando a cor do scrollbar da ajuda
    style.configure("round-info", troughcolor=COLOR_NIGHT_SKY, background=COLOR_TERRACOTTA) # Fundo do scrollbar e cor do polegar
    help_text_area['yscrollcommand'] = scrollbar_help.set
    
    scrollbar_help.pack(side="right", fill="y", padx=(0,5), pady=5)
    help_text_area.pack(fill="both", expand=True, padx=(5,0), pady=5)
    
    help_text_area.insert("1.0", texts["help_text"])
    help_text_area.config(state="disabled") # Torna o texto somente leitura

# Define um global para a animação de fundo
spots = []
def animate_organic_background():
    """Cria uma animação orgânica de fundo para o canvas principal."""
    global spots, canvas_background 
    try:
        # Verifica se o widget ainda existe e tem dimensões válidas
        if not root.winfo_exists() or not canvas_background.winfo_exists(): return 
        canvas_width = canvas_background.winfo_width()
        canvas_height = canvas_background.winfo_height()

        if canvas_width <= 1 or canvas_height <= 1: # Se a janela ainda não foi renderizada, espera
            root.after(50, animate_organic_background)
            return

        canvas_background.delete("all") # Limpa o canvas para redesenhar
        # Fundo principal do canvas como Noite Estrelada
        canvas_background.create_rectangle(0, 0, canvas_width, canvas_height, fill=COLOR_NIGHT_SKY, outline="") 

        # Cores mais orgânicas e terrosas para os "spots" (Terracota do Cerrado e tons relacionados)
        spot_colors = [COLOR_TERRACOTTA, "#D3A475", "#B8860B", "#A0522D", "#8B4513", "#F0A170"]

        # Inicializa ou re-inicializa os spots se a lista estiver vazia ou for menor que o esperado
        if not spots or len(spots) != 60: # Número de spots para a animação
            spots = []
            for _ in range(60): # Cria 60 spots
                x = random.randint(0, canvas_width)
                y = random.randint(0, canvas_height)
                size = random.randint(20, 80) # Tamanho aleatório dos spots
                color = random.choice(spot_colors)
                dx = random.uniform(-0.8, 0.8) # Velocidade horizontal
                dy = random.uniform(-0.8, 0.8) # Velocidade vertical
                spots.append({"x": x, "y": y, "size": size, "color": color, "dx": dx, "dy": dy})

        for spot in spots:
            spot["x"] += spot["dx"]
            spot["y"] += spot["dy"]
            
            # Colisão com as bordas
            if spot["x"] - spot["size"] < 0 or spot["x"] + spot["size"] > canvas_width:
                spot["dx"] *= -1 # Inverte a direção horizontal
                # Ajusta a posição para evitar que o spot "grude" na borda
                spot["x"] = max(spot["size"], min(spot["x"], canvas_width - spot["size"]))
            if spot["y"] - spot["size"] < 0 or spot["y"] + spot["size"] > canvas_height:
                spot["dy"] *= -1 # Inverte a direção vertical
                # Ajusta a posição
                spot["y"] = max(spot["size"], min(spot["y"], canvas_height - spot["size"]))
            
            canvas_background.create_oval(
                spot["x"] - spot["size"], spot["y"] - spot["size"],
                spot["x"] + spot["size"], spot["y"] + spot["size"],
                fill=spot["color"], outline="" # Sem borda para um visual mais suave
            )
        root.after(30, animate_organic_background) # Chama a função novamente após 30ms para animação
    except tk.TclError: # Captura erro se a janela for fechada durante a animação
        logging.info("Janela destruída, parando animação de fundo.")
    except Exception as e:
        logging.error(f"Erro na animação de fundo: {e}")

def animate_progress_bar():
    """Anima o estilo da barra de progresso."""
    try:
        if progress_bar.winfo_exists(): # Verifica se a barra de progresso existe
            current_value = progress_var.get()
            if current_value < progress_bar["maximum"]:
                progress_bar.configure(bootstyle="info-striped") # Estilo de progresso
            else:
                progress_bar.configure(bootstyle="success-striped") # Estilo de sucesso
    except tk.TclError:
        pass # Ignora se o widget já foi destruído

def animate_logo_pulse():
    """Faz o logo pulsar sutilmente."""
    try:
        if 'logo_label' in globals() and logo_label.winfo_exists(): # Verifica se o logo existe
            current_style = logo_label.cget("style")
            if current_style == "GuaraLogo.TLabel": # Alterna entre dois estilos para criar o pulso
                logo_label.configure(style="GuaraLogoPulse.TLabel") 
            else:
                logo_label.configure(style="GuaraLogo.TLabel")
            root.after(1000, animate_logo_pulse) # Chama novamente após 1 segundo
    except tk.TclError: # Captura erro se a janela for fechada
        logging.info("Janela destruída, parando animação de fundo.")
    except NameError: # Caso logo_label não tenha sido definido (e.g., falha no carregamento da imagem)
        logging.warning("logo_label não definido para animate_logo_pulse.")


def show_welcome_panel(): 
    """Exibe o painel de boas-vindas no frame de opções."""
    clear_options_frame()
    # Fundo do welcome frame em Luz Pura da Lua
    welcome_frame = ttkb.Frame(options_frame, background=COLOR_MOON_LIGHT)
    welcome_frame.pack(expand=True, fill="both")
    # Texto em Dark Earth (marrom escuro)
    welcome_label = ttkb.Label(
        welcome_frame,
        text="Bem-vindo ao Guará Codex!\nEscolha uma função nos painéis.", 
        font=("Papyrus", 20, "bold"), 
        foreground=COLOR_DARK_EARTH, 
        justify="center",
        anchor="center",
        background=COLOR_MOON_LIGHT # Garante que o fundo do label seja o mesmo do frame
    )
    welcome_label.place(relx=0.5, rely=0.5, anchor="center") # Centraliza o texto no frame

# --- Configuração da Interface Gráfica Principal (UI) ---
# Usando um tema base escuro como 'superhero' que já tem um bom contraste padrão
root = ttkb.Window(themename="superhero")
root.title(texts["title"])
root.state("zoomed") # Inicia maximizado
root.minsize(1100, 750) # Tamanho mínimo para garantir a legibilidade

style = ttkb.Style()

# Definindo as cores da Toca para fácil referência
COLOR_TERRACOTTA = '#e67e22' # Terracota do Cerrado (Primária da Alcateia)
COLOR_NIGHT_SKY = '#2c2c2c' # Noite Estrelada (Fundo Escuro Principal para contraste)
COLOR_MOON_LIGHT = '#ffffff' # Luz Pura da Lua (Fundo Claro Principal / Textos CLAROS)
COLOR_DARK_EARTH = '#4a3726' # Um marrom escuro para textos e detalhes, baseado no contraste.
COLOR_ACCENT_GREEN = '#6B8E23' # Um verde musgo ou cerrado para "sucesso" ou "info" alternativo
COLOR_LIGHT_GRAY_BG = '#F0F0F0' # Um cinza bem leve para fundos que precisam ser muito claros, mas não branco puro.

# Configurando estilos customizados com a paleta da Toca

# --- ESTILOS DE FRAMES E PAINÉIS ---
# Fundo padrão para TFrame é Noite Estrelada
style.configure("TFrame", background=COLOR_NIGHT_SKY)
# Frame com fundo claro (Luz Pura da Lua)
style.configure("GuaraFrameLight.TFrame", background=COLOR_MOON_LIGHT)
# Painéis de categoria com fundo claro e borda/título Terracota
style.configure("TLabelframe", background=COLOR_MOON_LIGHT, bordercolor=COLOR_TERRACOTTA)
style.configure("TLabelframe.Label", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA, font=("Verdana", 10, "bold"))
# Painel de Boas Vindas
style.configure("GuaraWelcome.TFrame", background=COLOR_MOON_LIGHT)

# --- ESTILOS DE BOTÕES ---
style.configure("TButton", font=("Verdana", 10), padding=8)

# Botões de Categoria (Manipulação e Conversão)
style.configure("GuaraOutlineCategory.TButton", foreground=COLOR_NIGHT_SKY, background=COLOR_MOON_LIGHT, bordercolor=COLOR_TERRACOTTA, borderwidth=2)
style.map("GuaraOutlineCategory.TButton",
          foreground=[('active', COLOR_MOON_LIGHT), ('!active', COLOR_NIGHT_SKY)],
          background=[('active', COLOR_TERRACOTTA), ('!disabled', COLOR_MOON_LIGHT)],
          bordercolor=[('active', COLOR_NIGHT_SKY)])

# Botões para compactação com níveis
style.configure("GuaraOutlineSuccessCategory.TButton", foreground=COLOR_NIGHT_SKY, background=COLOR_MOON_LIGHT, bordercolor=COLOR_ACCENT_GREEN, borderwidth=2)
style.map("GuaraOutlineSuccessCategory.TButton", background=[('active', COLOR_ACCENT_GREEN)], foreground=[('active', COLOR_MOON_LIGHT)])
style.configure("GuaraOutlineWarningCategory.TButton", foreground=COLOR_NIGHT_SKY, background=COLOR_MOON_LIGHT, bordercolor=COLOR_TERRACOTTA, borderwidth=2)
style.map("GuaraOutlineWarningCategory.TButton", background=[('active', COLOR_TERRACOTTA)], foreground=[('active', COLOR_MOON_LIGHT)])
style.configure("GuaraOutlineDangerCategory.TButton", foreground=COLOR_NIGHT_SKY, background=COLOR_MOON_LIGHT, bordercolor=COLOR_DARK_EARTH, borderwidth=2)
style.map("GuaraOutlineDangerCategory.TButton", background=[('active', COLOR_DARK_EARTH)], foreground=[('active', COLOR_MOON_LIGHT)])

# Botões do Rodapé (Ajuda/Sair)
style.configure('GuaraLinkButton.TButton', borderwidth=0, background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA, font=("Verdana", 10, "bold"))
style.map('GuaraLinkButton.TButton', foreground=[('active', COLOR_DARK_EARTH)])
style.configure('GuaraExitButton.TButton', borderwidth=0, background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, font=("Verdana", 10, "bold"))
style.map('GuaraExitButton.TButton', foreground=[('active', COLOR_TERRACOTTA)])


# --- ESTILOS DE LABELS (TEXTOS) ---
# Título e Logo
style.configure("GuaraLogo.TLabel", background=COLOR_MOON_LIGHT)
style.configure("GuaraHeaderTitle.TLabel", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH)
style.configure("GuaraHeaderSubtitle.TLabel", background=COLOR_MOON_LIGHT, foreground=COLOR_TERRACOTTA)
# Labels de descrição dentro dos painéis
style.configure("GuaraLabelDarkSmall.TLabel", background=COLOR_MOON_LIGHT, foreground=COLOR_NIGHT_SKY)
style.configure("GuaraLabelDark.TLabel", background=COLOR_MOON_LIGHT, foreground=COLOR_DARK_EARTH, font=("Verdana", 9))


# --- ESTILOS DE OUTROS WIDGETS ---
# Separador
style.configure("Guara.TSeparator", background=COLOR_TERRACOTTA)
# Barra de Progresso
style.configure("GuaraProgressBar.Horizontal.TProgressbar", troughcolor=COLOR_NIGHT_SKY, background=COLOR_TERRACOTTA, bordercolor=COLOR_TERRACOTTA)


# --- ESTILOS PARA JANELA DE PRÉVIA (mantidos para consistência se chamados) ---
style.configure("PreviewDark.TFrame", background=COLOR_NIGHT_SKY)
style.configure('PreviewInfo.TLabelframe', background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA, bordercolor=COLOR_TERRACOTTA)
style.configure('PreviewInfo.TLabelframe.Label', background=COLOR_NIGHT_SKY, foreground=COLOR_TERRACOTTA)

# Fundo principal do canvas como Noite Estrelada (moldura escura da aplicação)
canvas_background = tk.Canvas(root, highlightthickness=0, bg=COLOR_NIGHT_SKY)
canvas_background.pack(expand=True, fill="both")

# Frame para o conteúdo principal, centralizado e com tamanho fixo para estabilidade
content_outer_frame = ttkb.Frame(canvas_background, padding=20, style='TFrame')
content_outer_frame.place(relx=0.5, rely=0.5, anchor="center", width=1080, height=700)

# Content_frame agora é um TFrame com fundo claro para agrupar tudo
content_frame = ttkb.Frame(content_outer_frame, style='GuaraFrameLight.TFrame', padding=15)
content_frame.pack(expand=True, fill="both")


# Frame do cabeçalho - Fundo em Luz Pura da Lua
header_frame = ttkb.Frame(content_frame, style='GuaraFrameLight.TFrame')
header_frame.pack(fill="x", pady=(0, 10))
header_inner_frame = ttkb.Frame(header_frame, style='GuaraFrameLight.TFrame')
header_inner_frame.pack(anchor="center", pady=5)

# Carregamento do logo
try:
    logo_image_pil = Image.open("logo.png")
    logo_image_pil = logo_image_pil.resize((90, 90), Image.Resampling.LANCZOS)
    logo_photo = ImageTk.PhotoImage(logo_image_pil)
    logo_label = ttkb.Label(header_inner_frame, image=logo_photo, style="GuaraLogo.TLabel")
    logo_label.pack(pady=(5, 5))
    logo_label.image = logo_photo
except FileNotFoundError:
    logging.error("logo.png não encontrado! Usando fallback de texto.")
    logo_label = ttkb.Label(header_inner_frame, text="GC", font=("Georgia", 36, "bold"), style="GuaraLogo.TLabel")
    logo_label.pack(pady=(5, 5))
except Exception as e:
    logging.error(f"Erro ao carregar logo: {e}. Usando fallback de texto.")
    logo_label = ttkb.Label(header_inner_frame, text="GC", font=("Georgia", 36, "bold"), style="GuaraLogo.TLabel")
    logo_label.pack(pady=(5, 5))

# Título principal em Dark Earth, subtítulo em Terracota, ambos em fundo Luz Pura da Lua
title_main_label = ttkb.Label(header_inner_frame, text=texts["title"], font=("Trajan Pro", 26, "bold"), style='GuaraHeaderTitle.TLabel')
title_main_label.pack(pady=(0,2))
subtitle_main_label = ttkb.Label(header_inner_frame, text="Sua Alcateia de Soluções PDF", font=("Verdana", 11, "italic"), style='GuaraHeaderSubtitle.TLabel')
subtitle_main_label.pack(pady=(0, 5))

# Separador com cor Terracota
# CORREÇÃO APLICADA: Usando 'style' em vez de 'background'.
divider_main = ttkb.Separator(content_frame, orient=HORIZONTAL, style='Guara.TSeparator')
divider_main.pack(fill="x", pady=10, padx=20)


# Frame para o corpo interno (painéis de botões e opções) - Fundo em Luz Pura da Lua
inner_body_frame = ttkb.Frame(content_frame, style='GuaraFrameLight.TFrame')
inner_body_frame.pack(expand=True, fill="both", padx=5, pady=5)


# Frame para Status e Barra de Progresso - Fundo em Luz Pura da Lua
status_progress_frame = ttkb.Frame(inner_body_frame, style='GuaraFrameLight.TFrame')
status_progress_frame.pack(side=BOTTOM, fill=X, pady=(10,0), padx=5)

# Texto da barra de status em Noite Estrelada em fundo Luz Pura da Lua
status_label = ttkb.Label(status_progress_frame, text="", font=("Verdana", 9), anchor='w', style='GuaraLabelDarkSmall.TLabel')
status_label.pack(side=LEFT, padx=(5, 0), fill=X, expand=True)
progress_var = tk.DoubleVar()
# Barra de progresso com trilho Noite Estrelada e progresso em Terracota
progress_bar = ttkb.Progressbar(status_progress_frame, variable=progress_var, maximum=100, mode='determinate', length=250, style='GuaraProgressBar.Horizontal.TProgressbar')
progress_bar.pack(side=RIGHT, padx=(0,5))


# Painéis para os botões de funções - Fundo em Luz Pura da Lua
panels_main_frame = ttkb.Frame(inner_body_frame, style='GuaraFrameLight.TFrame')
panels_main_frame.pack(fill=BOTH, expand=True, pady=5)
panels_main_frame.grid_columnconfigure(0, weight=1)
panels_main_frame.grid_columnconfigure(1, weight=2)
panels_main_frame.grid_columnconfigure(2, weight=1)
panels_main_frame.grid_rowconfigure(0, weight=1)

# Painel de Manipulação de Páginas - Fundo Luz Pura da Lua, título/borda Terracota
manip_frame = ttkb.LabelFrame(panels_main_frame, text=texts["manipulation_frame"], padding=10, style='TLabelframe')
manip_frame.grid(row=0, column=0, padx=(0,5), pady=5, sticky="nsew")

# Descrição do painel de manipulação - Fundo Luz Pura da Lua, texto Noite Estrelada
desc_manip_label = ttkb.Label(manip_frame, text="Edite estrutura e páginas.", font=("Verdana", 8, "italic"), anchor="center", style='GuaraLabelDark.TLabel')
desc_manip_label.pack(pady=(0,10), fill=X)


# Botões de Manipulação (usando estilo GuaraOutlineCategory.TButton)
buttons_manip_config = [
    (texts["split_pdf"], on_split_button, texts["tooltip_split"], "split_pdf"),
    (texts["merge_pdf"], on_merge_button, texts["tooltip_merge"], "merge_pdf"),
    (texts["exclude_pages"], on_exclude_button, texts["tooltip_exclude"], "exclude_pages"),
    (texts["select_pages"], select_pages, texts["tooltip_select"], "select_pages"),
    (texts["add_pages"], add_selected_pages, texts["tooltip_add"], "add_pages"),
    (texts["replace_pages"], on_replace_button, texts["tooltip_replace"], "replace_pages")
]
for btn_text, btn_cmd, btn_tip, btn_key in buttons_manip_config:
    button = ttkb.Button(manip_frame, text=btn_text, command=btn_cmd, width=28, style='GuaraOutlineCategory.TButton')
    button.pack(pady=4, fill="x", padx=5)
    create_tooltip(button, btn_tip)

# Frame Central de Opções/Boas-Vindas - Fundo Luz Pura da Lua, título/borda Terracota
options_frame = ttkb.LabelFrame(panels_main_frame, text="Painel de Boas-Vindas / Opções", padding=15, style='TLabelframe')
options_frame.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")

# Painel de Conversão e Otimização - Fundo Luz Pura da Lua, título/borda Terracota
conv_frame = ttkb.LabelFrame(panels_main_frame, text=texts["conversion_frame"], padding=10, style='TLabelframe')
conv_frame.grid(row=0, column=2, padx=(5,0), pady=5, sticky="nsew")

# Descrição do painel de conversão - Fundo Luz Pura da Lua, texto Noite Estrelada
desc_conv_label = ttkb.Label(conv_frame, text="Transforme e otimize seus PDFs.", font=("Verdana", 8, "italic"), anchor="center", style='GuaraLabelDark.TLabel')
desc_conv_label.pack(pady=(0,10), fill=X)


# Botões de Conversão (usando estilo GuaraOutlineCategory.TButton para consistência)
buttons_conv_config = [
    (texts["compress_pdf"], compact_pdf, texts["tooltip_compress"], "compress_pdf"),
    (texts["anonymize_pdf"], anonymize_pdf, texts["tooltip_anonymize"], "anonymize_pdf"),
    (texts["convert_pdf"], on_convert_button, texts["tooltip_convert"], "convert_pdf"),
    (texts["pdf_jpg"], choose_conversion, texts["tooltip_pdf_jpg"], "pdf_jpg")
]
for btn_text, btn_cmd, btn_tip, btn_key in buttons_conv_config:
    button = ttkb.Button(conv_frame, text=btn_text, command=btn_cmd, width=28, style='GuaraOutlineCategory.TButton')
    button.pack(pady=4, fill="x", padx=5)
    create_tooltip(button, btn_tip)

# Frame do rodapé - Fundo em Luz Pura da Lua
footer_frame = ttkb.Frame(content_frame, style='GuaraFrameLight.TFrame')
footer_frame.pack(fill="x", side="bottom", pady=(10,0), padx=10)

# Separador do rodapé com cor Terracota
# CORREÇÃO APLICADA: Usando 'style' em vez de 'background'.
footer_divider = ttkb.Separator(footer_frame, orient=HORIZONTAL, style='Guara.TSeparator')
footer_divider.pack(fill="x", pady=(5,2))

footer_button_frame = ttkb.Frame(footer_frame, style='GuaraFrameLight.TFrame')
footer_button_frame.pack(pady=2)

# Botões de Ajuda e Sair (links) - Texto Terracota/Dark Earth, fundo Luz Pura da Lua
help_btn = ttkb.Button(footer_button_frame, text=texts["help"], command=show_help, width=15, style='GuaraLinkButton.TButton')
help_btn.pack(side="left", padx=20)
create_tooltip(help_btn, texts["tooltip_help"])

exit_btn = ttkb.Button(footer_button_frame, text=texts["exit"], command=root.quit, width=15, style='GuaraExitButton.TButton')
exit_btn.pack(side="left", padx=20)
create_tooltip(exit_btn, texts["tooltip_exit"])

footer_text_frame = ttkb.Frame(footer_frame, style='GuaraFrameLight.TFrame')
footer_text_frame.pack(pady=(2,5))

# Textos do rodapé em Noite Estrelada em fundo Luz Pura da Lua
footer_label_guaracodex = ttkb.Label(footer_text_frame, text=f"🐺 {texts['title']}", font=("Georgia", 9, "italic"), style='GuaraLabelDarkSmall.TLabel')
footer_label_guaracodex.pack(side="left", padx=5)
footer_label_dev = ttkb.Label(footer_text_frame, text=f"| {texts['footer']}", font=("Georgia", 9, "italic"), style='GuaraLabelDarkSmall.TLabel')
footer_label_dev.pack(side="left", padx=5)


# --- Inicialização da UI ---
if __name__ == "__main__":
    def start_animation_if_ready():
        """Inicia as animações de fundo e do logo quando a UI estiver renderizada."""
        try:
            if 'canvas_background' in globals() and canvas_background.winfo_exists() and canvas_background.winfo_width() > 1 and canvas_background.winfo_height() > 1 :
                animate_organic_background() # A animação de fundo foi removida para simplificar o layout e focar na funcionalidade
                if 'logo_label' in globals() and logo_label.winfo_exists():
                    # animate_logo_pulse() # A animação do logo foi removida para um look mais limpo
                    pass
            else:
                root.after(100, start_animation_if_ready)
        except tk.TclError:
            logging.info("Janela principal provavelmente fechada durante start_animation_if_ready.")
        except NameError as ne:
            logging.error(f"NameError em start_animation_if_ready: {ne}. Widget pode nao estar definido.")

    root.after(50, start_animation_if_ready)

    show_welcome_panel()

    root.bind("<Escape>", lambda event: root.quit())

    logging.info("Guará Codex iniciado com sucesso.")
    root.mainloop()
    logging.info("Guará Codex encerrado.")
